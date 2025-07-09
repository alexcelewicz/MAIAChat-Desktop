# rag_handler.py V3.0
import codecs
import os
import json
import hashlib
import logging
import math
import random
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union, Set
import faiss
import numpy as np

from config_manager import ConfigManager
from dataclasses import dataclass, field
from enum import Enum
import pdfplumber
import pandas as pd
from sentence_transformers import SentenceTransformer
from bs4 import BeautifulSoup
import nltk
from nltk.tokenize import sent_tokenize
import unicodedata
import re
from tqdm import tqdm
import chardet
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import pickle
import docx
import openpyxl
import tempfile
from rank_bm25 import BM25Okapi
from transformers import AutoTokenizer, AutoModel
import torch
from collections import Counter
from itertools import islice
import warnings
import html2text
import fitz  # PyMuPDF for better PDF handling
from io import BytesIO
from datetime import datetime
import difflib
import hashlib
from typing import Set, Callable, Iterator
import zipfile
import io
from functools import lru_cache
import threading
from collections import defaultdict

# Suppress specific warnings
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")
warnings.filterwarnings("ignore", category=UserWarning, module="torch")

# Set up logging configuration
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("rag_handler.log", encoding='utf-8')
    ]
)

logger = logging.getLogger("RAGHandler")

class EmbeddingProvider(Enum):
    OPENAI = "openai"
    SENTENCE_TRANSFORMER = "sentence_transformer"
    HUGGINGFACE = "huggingface"
    GEMINI = "gemini"
    COHERE = "cohere"

class ChunkingStrategy(Enum):
    FIXED_SIZE = "fixed_size"
    SEMANTIC = "semantic"
    RECURSIVE = "recursive"
    SLIDING_WINDOW = "sliding_window"
    CONTEXTUAL = "contextual"  # New strategy with better context preservation

@dataclass
class ChunkMetadata:
    file_name: str
    page_num: int
    chunk_index: int
    is_table: bool
    timestamp: int
    source_type: str  # e.g., 'pdf', 'excel', 'word', etc.
    sheet_name: Optional[str] = None  # For Excel files
    section_title: Optional[str] = None  # Section heading
    table_data: Optional[Dict] = None  # For tables
    row_data: Optional[Dict] = None  # For spreadsheet rows
    characters: Optional[Tuple[int, int]] = None  # Character positions in source
    keywords: List[str] = field(default_factory=list)  # Important keywords
    last_modified: Optional[int] = None  # File last modified timestamp
    embedding_model: Optional[str] = None  # Which model created the embedding
    created_at: int = field(default_factory=lambda: int(time.time()))
    version: Optional[str] = None  # Document version tracking
    previous_versions: List[str] = field(default_factory=list)  # Previous hashes
    context_before: Optional[str] = None  # Context preceding this chunk
    context_after: Optional[str] = None  # Context following this chunk
    importance_score: float = 1.0  # Score indicating chunk importance
    language: Optional[str] = None  # Detected language of the chunk

# Initialize NLTK resources with better error handling
def initialize_nltk():
    """Initialize NLTK, strongly preferring locally managed resources."""
    try:
        # Determine the correct path to the 'nltk_data' directory relative to this file
        # Assumes rag_handler.py is in the project root or a subdirectory.
        # If rag_handler.py is in project_root/sub_dir/rag_handler.py
        # then nltk_data should be in project_root/nltk_data
        base_dir = Path(__file__).resolve().parent.parent # Adjust if rag_handler.py is deeper
        nltk_data_dir = base_dir / 'nltk_data'
        
        flag_file = nltk_data_dir / 'nltk_download_complete.flag'

        if flag_file.exists() and nltk_data_dir.is_dir():
            if str(nltk_data_dir) not in nltk.data.path: # Add to path if not already there
                nltk.data.path.insert(0, str(nltk_data_dir)) # Prepend for priority
            logger.info(f"NLTK using local resources from: {nltk_data_dir}")
            # Optional: Verify key packages are loadable from this path
            try:
                nltk.data.find('tokenizers/punkt')
                nltk.data.find('corpora/wordnet') # Example check
                logger.info("Key NLTK resources confirmed in local path.")
            except LookupError as le:
                logger.warning(f"NLTK flag file exists, but key resource '{le.args[0]}' not found in {nltk_data_dir}. "
                               "Consider re-running 'python install_dependencies.py'.")
        else:
            logger.error(
                f"Local NLTK data setup incomplete (flag: {flag_file.exists()}, dir: {nltk_data_dir.is_dir()}). "
                "Please run 'python install_dependencies.py' for optimal setup. "
                "Application will attempt to use system NLTK paths, which may be unreliable."
            )
            # No nltk.download() calls here. The user must run the setup script.
        
        # Test tokenization to confirm NLTK is functional
        from nltk.tokenize import sent_tokenize
        sent_tokenize("This is a test.")
        logger.info("NLTK tokenization test successful after initialization.")

    except ImportError:
        logger.critical("NLTK library not found. Please install it by running 'pip install -r requirements.txt'.")
        raise
    except Exception as e:
        logger.error(f"Critical error initializing NLTK: {e}. RAG functionality will be impaired.")
        # Depending on desired behavior, you might re-raise or allow graceful degradation.

# Initialize NLTK at module level
initialize_nltk()

def get_stopwords():
    """Get stopwords with fallback mechanism."""
    try:
        return nltk.corpus.stopwords.words('english')
    except:
        return {
            'a', 'an', 'the', 'and', 'or', 'but', 'if', 'because', 'as', 'what',
            'when', 'where', 'how', 'which', 'who', 'whom', 'this', 'that', 'these',
            'those', 'to', 'of', 'in', 'for', 'with', 'by', 'at', 'on', 'about',
            'against', 'between', 'into', 'through', 'during', 'before', 'after',
            'above', 'below', 'from', 'up', 'down', 'is', 'am', 'are', 'was',
            'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do',
            'does', 'did', 'doing', 'would', 'should', 'could', 'ought', 'i',
            'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us',
            'them', 'my', 'your', 'his', 'hers', 'its', 'ours', 'theirs', 'myself',
            'yourself', 'himself', 'herself', 'itself', 'ourselves', 'yourselves',
            'themselves'
        }

class TextPreprocessor:
    """Advanced text preprocessing for better chunking and embedding."""

    def __init__(self):
        self.stopwords = get_stopwords()
        self.h2t = html2text.HTML2Text()
        self.h2t.ignore_links = False
        self.h2t.ignore_images = False
        self.h2t.ignore_tables = False
        self._keyword_cache = {}  # Cache for keyword extraction
        self._language_cache = {}  # Cache for language detection
        self._cache_lock = threading.Lock()  # Thread-safe cache access

    def clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        if not text:
            return ""

        # Replace multiple newlines with single newline
        text = re.sub(r'\n{2,}', '\n', text)

        # Normalize unicode characters
        text = unicodedata.normalize('NFKC', text)

        # Replace tabs with spaces
        text = text.replace('\t', ' ')

        # Remove control characters
        text = ''.join(ch for ch in text if unicodedata.category(ch)[0] != 'C' or ch == '\n')

        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)

        return text.strip()

    def html_to_text(self, html_content: str) -> str:
        """Convert HTML to markdown text."""
        return self.h2t.handle(html_content)

    def extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """Extract important keywords from text using improved techniques."""
        # Check cache first
        text_hash = hashlib.md5(text.encode()).hexdigest()
        with self._cache_lock:
            if text_hash in self._keyword_cache:
                return self._keyword_cache[text_hash][:max_keywords]

        # Tokenize and lowercase
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())

        # Remove stopwords
        words = [w for w in words if w not in self.stopwords]

        if not words:
            return []

        # Get term frequency
        word_counts = Counter(words)

        # Calculate TF-IDF-like weighting - give higher weight to less common terms
        # by penalizing very frequent terms that appear many times in the document
        total_words = len(words)
        word_weights = {}

        for word, count in word_counts.items():
            # Simple TF (term frequency)
            tf = count / total_words

            # Discount for very common words (occurs >5 times and >10% of text)
            repeated_word_discount = 1.0
            if count > 5 and tf > 0.1:
                repeated_word_discount = 0.5

            # Boost for words in title case (often important entities)
            title_case_boost = 1.5 if word.title() in text else 1.0

            # Boost for longer words (often more meaningful)
            length_boost = min(1.0 + (len(word) / 20), 1.5)

            # Combine factors
            word_weights[word] = tf * repeated_word_discount * title_case_boost * length_boost

        # Get top weighted words
        keywords = [word for word, _ in sorted(word_weights.items(),
                                              key=lambda x: x[1],
                                              reverse=True)[:max_keywords]]

        # Cache results
        with self._cache_lock:
            self._keyword_cache[text_hash] = keywords

            # Limit cache size
            if len(self._keyword_cache) > 10000:
                # Remove random 20% of entries when cache gets too large
                keys_to_remove = list(self._keyword_cache.keys())[:2000]
                for k in keys_to_remove:
                    self._keyword_cache.pop(k, None)

        return keywords

    def detect_language(self, text: str) -> str:
        """Detect the language of the text with caching."""
        if not text or len(text.strip()) < 20:
            return "en"  # Default for very short texts

        # Check cache first
        text_hash = hashlib.md5(text[:1000].encode()).hexdigest()  # Use just first 1000 chars for hash
        with self._cache_lock:
            if text_hash in self._language_cache:
                return self._language_cache[text_hash]

        # Multiple detection methods for robustness
        result = "en"  # Default to English

        # Try method 1: langdetect
        try:
            from langdetect import detect
            result = detect(text)
        except Exception:
            # Try method 2: simple character-based detection
            try:
                # Basic heuristic language detection as fallback
                # Count character frequencies typical for certain languages
                cyrillic = sum(1 for c in text if '\u0400' <= c <= '\u04FF')
                cjk = sum(1 for c in text if '\u4E00' <= c <= '\u9FFF')
                arabic = sum(1 for c in text if '\u0600' <= c <= '\u06FF')

                text_len = len(text)

                if cyrillic / text_len > 0.3:
                    result = "ru"  # Russian or other Cyrillic-based
                elif cjk / text_len > 0.3:
                    result = "zh"  # Chinese or other CJK
                elif arabic / text_len > 0.3:
                    result = "ar"  # Arabic
            except Exception:
                # Default to English if everything fails
                result = "en"

        # Cache the result
        with self._cache_lock:
            self._language_cache[text_hash] = result
            # Limit cache size
            if len(self._language_cache) > 10000:
                keys_to_remove = list(self._language_cache.keys())[:2000]
                for k in keys_to_remove:
                    self._language_cache.pop(k, None)

        return result

    def detect_section_titles(self, text: str) -> List[Tuple[int, str]]:
        """Detect section titles and their positions."""
        # Match common heading patterns
        heading_patterns = [
            r'^#{1,6}\s+(.+)$',  # Markdown headings
            r'^([A-Z][A-Za-z0-9\s]{2,50})$',  # Capitalized text line
            r'^(\d+\.\s+[A-Z][A-Za-z0-9\s]{2,50})$',  # Numbered headings
            r'^([IVX]+\.\s+[A-Z][A-Za-z0-9\s]{2,50})$',  # Roman numerals
        ]

        titles = []
        position = 0

        for line in text.split('\n'):
            line = line.strip()
            for pattern in heading_patterns:
                match = re.match(pattern, line)
                if match:
                    titles.append((position, match.group(1).strip()))
                    break
            position += len(line) + 1  # +1 for the newline

        return titles

    def segment_by_delimiter(self, text: str, delimiters: List[str]) -> List[str]:
        """Split text by multiple delimiters while preserving delimiters."""
        pattern = '|'.join(map(re.escape, delimiters))
        parts = re.split(f'({pattern})', text)

        result = []
        for i in range(0, len(parts), 2):
            if i + 1 < len(parts):
                result.append(parts[i] + parts[i+1])
            else:
                result.append(parts[i])

        return [r for r in result if r.strip()]

def tokenize_text(text: str) -> List[str]:
    """
    Tokenize text into sentences with robust fallback methods.
    """
    if not text or not text.strip():
        return []

    # First try with standard NLTK tokenizer
    try:
        return sent_tokenize(text)
    except Exception as e:
        logger.warning(f"NLTK tokenization failed: {e}")

        # Check for local NLTK data directory first
        nltk_data_dir = os.path.join(os.path.dirname(__file__), 'nltk_data')
        if os.path.exists(nltk_data_dir):
            # Add the local directory to NLTK's search path if not already added
            if nltk_data_dir not in nltk.data.path:
                nltk.data.path.append(nltk_data_dir)
        else:
            # Try to download punkt and punkt_tab if needed
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('punkt_tab', quiet=True)
            except Exception:
                pass

        # Second try with punkt directly
        try:
            from nltk.tokenize.punkt import PunktSentenceTokenizer
            tokenizer = PunktSentenceTokenizer()
            return tokenizer.tokenize(text)
        except Exception as e2:
            logger.warning(f"Secondary NLTK tokenization failed: {e2}")

        # Final fallback: split by common sentence terminators
        # This is a simple but reliable method that will always work
        basic_splits = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in basic_splits if s.strip()]

class RAGHandler:
    def __init__(self,
                persist_directory: str = "./knowledge_base",
                config_path: str = "config.json",
                use_openai: bool = True,  # For backward compatibility
                embedding_model: str = "text-embedding-3-small",
                st_model_name: str = "all-mpnet-base-v2",  # For backward compatibility
                dimension: int = 768,
                chunk_size: int = 500,
                chunk_overlap: int = 50,
                embedding_provider: Optional[Union[str, EmbeddingProvider]] = None,
                chunking_strategy: Union[str, ChunkingStrategy] = ChunkingStrategy.SEMANTIC,
                cache_dir: Optional[str] = None):
        """Initialize the RAGHandler with enhanced features."""
        self.persist_directory = Path(persist_directory)
        self.persist_directory.mkdir(parents=True, exist_ok=True)

        # Set up caching directory for models
        self.cache_dir = cache_dir or os.path.join(self.persist_directory, "model_cache")
        os.makedirs(self.cache_dir, exist_ok=True)

        # Load configuration
        try:
            self.config = self._load_config(config_path)
        except Exception as e:
            logger.error(f"Error loading config: {e}")
            self.config = {
                "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY"),
                "GEMINI_API_KEY": os.getenv("GEMINI_API_KEY"),
                "GROQ_API_KEY": os.getenv("GROQ_API_KEY"),
                "GROK_API_KEY": os.getenv("GROK_API_KEY"),
                "COHERE_API_KEY": os.getenv("COHERE_API_KEY"),
                "GOOGLE_API_KEY": os.getenv("GOOGLE_API_KEY"),
                "GOOGLE_SEARCH_ENGINE_ID": os.getenv("GOOGLE_SEARCH_ENGINE_ID")
            }

        # Instantiate ConfigManager to read current settings
        self.config_manager = ConfigManager()
        self.ultra_safe_mode = self.config_manager.get('RAG_ULTRA_SAFE_MODE', False)
        self.safe_retrieval_mode = self.config_manager.get('RAG_SAFE_RETRIEVAL_MODE', False)
        self.embedding_device_setting = self.config_manager.get('EMBEDDING_DEVICE', 'cpu')
        logger.info(f"RAGHandler: Ultra Safe Mode: {self.ultra_safe_mode}, Safe Retrieval Mode: {self.safe_retrieval_mode}, Embedding Device: {self.embedding_device_setting}")

        # Handle backward compatibility with use_openai parameter
        if embedding_provider is None:
            # Convert use_openai to the appropriate embedding provider
            self.embedding_provider = EmbeddingProvider.OPENAI if use_openai else EmbeddingProvider.SENTENCE_TRANSFORMER

            # Set model name based on provider for backward compatibility
            if use_openai:
                self.embedding_model_name = embedding_model
            else:
                self.embedding_model_name = st_model_name
        else:
            # Use the new embedding_provider parameter if provided
            if isinstance(embedding_provider, str):
                try:
                    self.embedding_provider = EmbeddingProvider(embedding_provider.lower())
                except ValueError:
                    logger.warning(f"Invalid embedding provider '{embedding_provider}', defaulting to SENTENCE_TRANSFORMER")
                    self.embedding_provider = EmbeddingProvider.SENTENCE_TRANSFORMER
            else:
                self.embedding_provider = embedding_provider

            # Use provided embedding_model_name
            self.embedding_model_name = embedding_model

        # Ensure correct model dimensions based on provider/model
        if self.embedding_provider == EmbeddingProvider.OPENAI:
            if self.embedding_model_name == "text-embedding-3-small":
                self.dimension = 1536
            elif self.embedding_model_name == "text-embedding-3-large":
                self.dimension = 3072
            else:
                self.dimension = dimension
        else:
            self.dimension = dimension

        # Initialize chunking settings
        if isinstance(chunking_strategy, str):
            try:
                self.chunking_strategy = ChunkingStrategy(chunking_strategy.lower())
            except ValueError:
                logger.warning(f"Invalid chunking strategy '{chunking_strategy}', defaulting to SEMANTIC")
                self.chunking_strategy = ChunkingStrategy.SEMANTIC
        else:
            self.chunking_strategy = chunking_strategy

        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Initialize text preprocessor
        self.text_processor = TextPreprocessor()

        # Initialize embedding model
        self._initialize_embedding_model()

        # Initialize FAISS index with customizable metric
        # Use IP (inner product) for normalized embeddings, L2 for others
        self.metric_type = faiss.METRIC_INNER_PRODUCT  # or faiss.METRIC_L2
        self.index = self._create_index()
        logger.info(f"Initialized FAISS index with dimension {self.dimension}")

        # Initialize storage for chunks and metadata
        self.chunks: List[str] = []
        self.metadatas: List[ChunkMetadata] = []

        # Initialize paths
        self.index_path = self.persist_directory / "faiss.index"
        self.metadata_path = self.persist_directory / "metadata.pkl"
        self.file_registry_path = self.persist_directory / "file_registry.json"
        self.embeddings_cache_path = self.persist_directory / "embeddings_cache.pkl"

        # Initialize embeddings cache
        self.embeddings_cache = self._load_embeddings_cache()

        # Load existing data
        self.load_index()
        self.file_registry = self._load_file_registry()

        # Check and repair file registry if needed
        self.check_and_repair_file_registry()

        # For multithreaded processing
        self.max_threads = min(os.cpu_count() or 4, 8)  # Limit max threads

        logger.info(f"RAGHandler initialized with {self.embedding_provider.value} using model {self.embedding_model_name}")

    def _initialize_embedding_model(self):
        """Initialize the embedding model based on provider with improved fallback logic."""
        try:
            if self.embedding_provider == EmbeddingProvider.OPENAI:
                try:
                    # Import here to avoid requiring openai when not using it
                    from openai import OpenAI
                    self.embedding_client = OpenAI(api_key=self.config.get('OPENAI_API_KEY'))
                    logger.info("Initialized OpenAI embedding client")

                    # Set dimension based on model
                    if self.embedding_model_name == "text-embedding-3-small":
                        self.dimension = 1536
                    elif self.embedding_model_name == "text-embedding-3-large":
                        self.dimension = 3072
                    else:
                        self.dimension = 1536  # Default for OpenAI embeddings
                except ImportError:
                    logger.warning("OpenAI package not found. Falling back to Sentence Transformer.")
                    self.embedding_provider = EmbeddingProvider.SENTENCE_TRANSFORMER
                    self.embedding_model_name = "all-mpnet-base-v2"
                    self._initialize_embedding_model()
                    return

            elif self.embedding_provider == EmbeddingProvider.SENTENCE_TRANSFORMER:
                try:
                    import torch
                    device = self.embedding_device_setting
                    if device == 'auto':
                        device = 'cuda' if torch.cuda.is_available() else 'cpu'
                        logger.info(f"Using {device.upper()} for SentenceTransformer (auto-detected)")
                    else:
                        logger.info(f"Using {device.upper()} for SentenceTransformer (from config)")

                    self.st_model = SentenceTransformer(
                        self.embedding_model_name,
                        cache_folder=self.cache_dir,
                        device=device
                    )
                    self.dimension = self.st_model.get_sentence_embedding_dimension()
                    logger.info(f"Initialized Sentence Transformer model: '{self.embedding_model_name}' with dimension {self.dimension}")
                except Exception as e:
                    logger.warning(f"Failed to load primary Sentence Transformer model '{self.embedding_model_name}': {e}. Attempting fallback to 'all-MiniLM-L6-v2'.")
                    self.embedding_model_name = "all-MiniLM-L6-v2"  # Set fallback model
                    self.dimension = 384  # Update dimension for this specific fallback model
                    try:
                        import torch
                        device = self.embedding_device_setting
                        if device == 'auto':
                            device = 'cuda' if torch.cuda.is_available() else 'cpu'
                            logger.info(f"Using {device.upper()} for SentenceTransformer (auto-detected fallback)")
                        else:
                            logger.info(f"Using {device.upper()} for SentenceTransformer (from config for fallback)")

                        self.st_model = SentenceTransformer(
                            self.embedding_model_name,
                            cache_folder=self.cache_dir,
                            device=device
                        )
                        self.dimension = self.st_model.get_sentence_embedding_dimension()
                        logger.info(f"Successfully initialized fallback Sentence Transformer model: '{self.embedding_model_name}' with dimension {self.dimension}")
                    except Exception as fallback_error:
                        logger.critical(f"CRITICAL ERROR: Failed to load even the fallback model 'all-MiniLM-L6-v2': {fallback_error}. RAG functionality will be severely impaired or non-functional. Please install a compatible model or troubleshoot your environment.")
                        logger.critical(traceback.format_exc())
                        self.st_model = None  # Set to None to indicate failure
                        self.dimension = 768  # Default to a common dimension, but embeddings will be zero
                        raise

            elif self.embedding_provider == EmbeddingProvider.HUGGINGFACE:
                try:
                    self.tokenizer = AutoTokenizer.from_pretrained(
                        self.embedding_model_name,
                        cache_dir=self.cache_dir
                    )
                    self.hf_model = AutoModel.from_pretrained(
                        self.embedding_model_name,
                        cache_dir=self.cache_dir
                    )
                    # Move model to CUDA if available
                    if torch.cuda.is_available():
                        self.hf_model = self.hf_model.cuda()
                        logger.info("Using CUDA for HuggingFace model")
                    else:
                        logger.info("Using CPU for HuggingFace model")

                    # Get embedding dimension from model config
                    self.dimension = self.hf_model.config.hidden_size
                    logger.info(f"Initialized HuggingFace model: {self.embedding_model_name} with dimension {self.dimension}")
                except ImportError:
                    logger.warning("HuggingFace transformers package not found. Falling back to Sentence Transformer.")
                    self.embedding_provider = EmbeddingProvider.SENTENCE_TRANSFORMER
                    self.embedding_model_name = "all-mpnet-base-v2"
                    self._initialize_embedding_model()
                    return

            elif self.embedding_provider == EmbeddingProvider.GEMINI:
                try:
                    # Import here to avoid requiring google-generativeai when not using it
                    import google.generativeai as genai
                    genai.configure(api_key=self.config.get('GEMINI_API_KEY'))
                    self.gemini_client = genai
                    self.dimension = 768  # Default for Gemini embeddings
                    logger.info("Initialized Gemini embedding client")
                except ImportError:
                    logger.warning("google-generativeai package not found. Falling back to Sentence Transformer.")
                    self.embedding_provider = EmbeddingProvider.SENTENCE_TRANSFORMER
                    self.embedding_model_name = "all-mpnet-base-v2"
                    self._initialize_embedding_model()
                    return

            elif self.embedding_provider == EmbeddingProvider.COHERE:
                try:
                    # Import here to avoid requiring cohere when not using it
                    import cohere
                    self.cohere_client = cohere.Client(self.config.get('COHERE_API_KEY'))
                    self.dimension = 1024  # Default for Cohere embeddings
                    logger.info("Initialized Cohere embedding client")
                except ImportError:
                    logger.warning("cohere package not found. Falling back to Sentence Transformer.")
                    self.embedding_provider = EmbeddingProvider.SENTENCE_TRANSFORMER
                    self.embedding_model_name = "all-mpnet-base-v2"
                    self._initialize_embedding_model()
                    return

            else:
                logger.warning(f"Unsupported embedding provider: {self.embedding_provider}, falling back to Sentence Transformer")
                self.embedding_provider = EmbeddingProvider.SENTENCE_TRANSFORMER
                self.embedding_model_name = "all-mpnet-base-v2"
                self._initialize_embedding_model()

        except Exception as e:
            logger.error(f"Error initializing embedding model: {e}")
            logger.error(traceback.format_exc())
            logger.warning("Falling back to Sentence Transformer with all-mpnet-base-v2 on CPU")
            self.embedding_provider = EmbeddingProvider.SENTENCE_TRANSFORMER
            self.embedding_model_name = "all-mpnet-base-v2"
            try:
                # Force CPU usage to avoid MPS/GPU issues
                import torch
                device = 'cpu'  # Force CPU usage
                logger.info(f"Using device: {device} for fallback SentenceTransformer")

                self.st_model = SentenceTransformer(
                    self.embedding_model_name,
                    cache_folder=self.cache_dir,
                    device=device
                )
                self.dimension = self.st_model.get_sentence_embedding_dimension()
                logger.info(f"Initialized fallback Sentence Transformer model with dimension {self.dimension}")
            except Exception as fallback_error:
                logger.error(f"Critical error initializing fallback model: {fallback_error}")
                logger.error(traceback.format_exc())
                self.st_model = None
                self.dimension = 768
                raise

    def _check_embedding_model_compatibility(self) -> bool:
        """Check if the current embedding model is compatible with existing index and metadata."""
        try:
            if not self.index_path.exists():
                return True  # No existing index, so compatible
            
            # Check if we have metadata to compare
            if self.metadata_path.exists():
                try:
                    with open(self.metadata_path, 'rb') as f:
                        data = pickle.load(f)
                    
                    # Check model compatibility
                    if 'version' in data and data['version'] == '3.0':
                        loaded_model = data.get('embedding_model')
                        loaded_provider = data.get('embedding_provider')
                        loaded_dimension = data.get('dimension')
                        
                        if (loaded_model and loaded_model != self.embedding_model_name) or \
                           (loaded_provider and loaded_provider != self.embedding_provider.value) or \
                           (loaded_dimension and loaded_dimension != self.dimension):
                            logger.warning(f"Embedding model mismatch detected:")
                            logger.warning(f"  Loaded: {loaded_provider}/{loaded_model} (dim: {loaded_dimension})")
                            logger.warning(f"  Current: {self.embedding_provider.value}/{self.embedding_model_name} (dim: {self.dimension})")
                            return False
                    
                    return True
                except Exception as e:
                    logger.warning(f"Error checking metadata compatibility: {e}")
                    return False
            
            return True
        except Exception as e:
            logger.error(f"Error checking embedding model compatibility: {e}")
            return False

    def _handle_embedding_model_change(self) -> bool:
        """Handle embedding model changes by prompting for re-indexing or auto-handling."""
        try:
            logger.warning("Embedding model change detected. This requires re-indexing the knowledge base.")
            
            # Check if we have existing data
            if self.index_path.exists() and len(self.get_indexed_files()) > 0:
                logger.info(f"Found {len(self.get_indexed_files())} indexed files that need re-indexing.")
                
                # Create backup of existing index
                backup_dir = Path(self.persist_directory) / "backup"
                backup_dir.mkdir(exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                if self.index_path.exists():
                    backup_index_path = backup_dir / f"index_backup_{timestamp}.faiss"
                    import shutil
                    shutil.copy2(self.index_path, backup_index_path)
                    logger.info(f"Created backup of existing index: {backup_index_path}")
                
                if self.metadata_path.exists():
                    backup_metadata_path = backup_dir / f"metadata_backup_{timestamp}.pkl"
                    shutil.copy2(self.metadata_path, backup_metadata_path)
                    logger.info(f"Created backup of existing metadata: {backup_metadata_path}")
                
                # Clear existing index and metadata
                self.index = self._create_index()
                self.chunks = []
                self.metadatas = []
                self.file_registry = {}
                
                # Save empty state
                self.save_index()
                
                logger.info("Cleared existing index and metadata. Knowledge base needs to be re-indexed.")
                logger.info("Please re-add your files to the knowledge base to use the new embedding model.")
                
                return False  # Indicates re-indexing is needed
            else:
                logger.info("No existing indexed files found. New index will be created with current model.")
                return True  # No re-indexing needed
                
        except Exception as e:
            logger.error(f"Error handling embedding model change: {e}")
            return False

    def _create_index(self) -> faiss.Index:
        """Create appropriate FAISS index based on dimension and metric."""
        if self.metric_type == faiss.METRIC_INNER_PRODUCT:
            return faiss.IndexFlatIP(self.dimension)
        else:  # L2 distance
            return faiss.IndexFlatL2(self.dimension)

    def _get_embeddings(self, texts: List[str]) -> np.ndarray:
        """Get embeddings based on the selected provider."""
        if not texts:
            logger.warning("Empty texts provided for embedding")
            return np.array([])

        # Preprocess texts
        processed_texts = [self.text_processor.clean_text(text) for text in texts]

        # Check cache first for each text
        cache_hits = []
        texts_to_embed = []
        indices_to_embed = []

        for i, text in enumerate(processed_texts):
            text_hash = hashlib.md5(text.encode()).hexdigest()
            if text_hash in self.embeddings_cache:
                cache_hits.append((i, self.embeddings_cache[text_hash]))
            else:
                texts_to_embed.append(text)
                indices_to_embed.append(i)

        if not texts_to_embed:
            # All embeddings found in cache
            sorted_embeddings = [embedding for _, embedding in sorted(cache_hits, key=lambda x: x[0])]
            return np.array(sorted_embeddings)

        try:
            if self.embedding_provider == EmbeddingProvider.OPENAI:
                response = self.embedding_client.embeddings.create(
                    input=texts_to_embed,
                    model=self.embedding_model_name
                )
                new_embeddings = np.array([embedding.embedding for embedding in response.data])

            elif self.embedding_provider == EmbeddingProvider.SENTENCE_TRANSFORMER:
                if self.st_model is None:
                    logger.critical("CRITICAL: SentenceTransformer model (st_model) is not initialized in RAGHandler. Cannot generate embeddings. RAG will not function correctly.")
                    # Option 1: Raise an exception
                    # raise RuntimeError("RAGHandler: SentenceTransformer model is not available.")
                    # Option 2: Return an empty array and rely on downstream checks
                    return np.array([]) # Or np.zeros((len(processed_texts), self.dimension)) with a log
                try:
                    # Process in smaller batches to avoid memory issues
                    batch_size = 8  # Smaller batch size to avoid memory issues
                    all_batches = []

                    # Use tqdm for progress tracking
                    for i in tqdm(range(0, len(texts_to_embed), batch_size), desc="Batches"):
                        batch = texts_to_embed[i:i+batch_size]
                        # Enhanced batch processing with individual text fallback
                        try:
                            batch_embeddings = self.st_model.encode(batch, normalize_embeddings=True)
                            all_batches.append(batch_embeddings)
                        except Exception as batch_error:
                            logger.error(f"Error encoding batch {i//batch_size}: {batch_error}. Trying individual encoding for this batch.")
                            individual_embeddings_for_batch = []
                            for text_in_batch in batch:
                                try:
                                    # Ensure single text is passed as a list
                                    single_text_embedding = self.st_model.encode([text_in_batch], normalize_embeddings=True)
                                    individual_embeddings_for_batch.append(single_text_embedding[0]) # Get the single embedding
                                except Exception as single_text_error:
                                    logger.error(f"Error encoding single text: {single_text_error}. Using zero embedding.")
                                    individual_embeddings_for_batch.append(np.zeros(self.dimension))
                            
                            if individual_embeddings_for_batch:
                                all_batches.append(np.array(individual_embeddings_for_batch))
                            else: # Should not happen if batch was not empty
                                all_batches.append(np.zeros((len(batch), self.dimension)))

                    # Combine all batches
                    if all_batches:
                        new_embeddings = np.vstack(all_batches)
                    else:
                        new_embeddings = np.zeros((len(texts_to_embed), self.dimension))
                except Exception as e:
                    logger.error(f"Error in sentence transformer encoding: {e}")
                    logger.error(traceback.format_exc())
                    # Fallback to zero embeddings
                    logger.warning("Using zero embeddings as fallback")
                    new_embeddings = np.zeros((len(texts_to_embed), self.dimension))

            elif self.embedding_provider == EmbeddingProvider.HUGGINGFACE:
                new_embeddings = []
                batch_size = 8  # Process in batches to avoid OOM

                for i in range(0, len(texts_to_embed), batch_size):
                    batch_texts = texts_to_embed[i:i+batch_size]

                    # Tokenize
                    encoded_input = self.tokenizer(
                        batch_texts,
                        padding=True,
                        truncation=True,
                        max_length=512,
                        return_tensors='pt'
                    )

                    # Move to GPU if available
                    if torch.cuda.is_available():
                        encoded_input = {k: v.cuda() for k, v in encoded_input.items()}

                    # Get model output
                    with torch.no_grad():
                        model_output = self.hf_model(**encoded_input)

                    # Mean pooling of token embeddings
                    token_embeddings = model_output.last_hidden_state
                    attention_mask = encoded_input['attention_mask']

                    # Mask padded tokens
                    input_mask_expanded = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
                    sum_embeddings = torch.sum(token_embeddings * input_mask_expanded, 1)
                    sum_mask = torch.sum(input_mask_expanded, 1)

                    # Avoid division by zero
                    sum_mask = torch.clamp(sum_mask, min=1e-9)
                    batch_embeddings = sum_embeddings / sum_mask

                    # Move to CPU and convert to numpy
                    batch_embeddings = batch_embeddings.cpu().numpy()

                    # Normalize
                    batch_embeddings = batch_embeddings / np.linalg.norm(batch_embeddings, axis=1, keepdims=True)

                    new_embeddings.append(batch_embeddings)

                new_embeddings = np.vstack(new_embeddings)

            elif self.embedding_provider == EmbeddingProvider.GEMINI:
                results = []
                for text in texts_to_embed:
                    result = self.gemini_client.embed_content(
                        model="models/embedding-001",
                        content=text,
                        task_type="retrieval_document"
                    )
                    results.append(result["embedding"])
                new_embeddings = np.array(results)

            elif self.embedding_provider == EmbeddingProvider.COHERE:
                results = self.cohere_client.embed(
                    texts=texts_to_embed,
                    model="embed-english-v3.0",
                    input_type="search_document"
                )
                new_embeddings = np.array(results.embeddings)

            else:
                raise ValueError(f"Unsupported embedding provider: {self.embedding_provider}")

            # Update cache with new embeddings
            for idx, text in zip(indices_to_embed, texts_to_embed):
                text_hash = hashlib.md5(text.encode()).hexdigest()
                self.embeddings_cache[text_hash] = new_embeddings[indices_to_embed.index(idx)]

            # Save cache periodically (every 100 new embeddings)
            if len(indices_to_embed) > 100:
                self._save_embeddings_cache()

            # Combine cached and new embeddings
            all_embeddings = np.zeros((len(texts), self.dimension))
            for i, embedding in cache_hits:
                all_embeddings[i] = embedding
            for idx, embedding in zip(indices_to_embed, new_embeddings):
                all_embeddings[idx] = embedding
            return all_embeddings
        except Exception as e:
            logger.error(f"Error getting embeddings: {e}")
            logger.error(traceback.format_exc())
            logger.warning("Using zero embeddings as fallback for all texts")
            return np.zeros((len(texts), self.dimension))

    def _load_config(self, config_path: str) -> Dict[str, Any]:
        """Load configuration from config.json file."""
        try:
            if not os.path.exists(config_path):
                logger.warning(f"Config file not found at {config_path}")
                return {
                    "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY"),
                    "GEMINI_API_KEY": os.getenv("GEMINI_API_KEY"),
                    "GROQ_API_KEY": os.getenv("GROQ_API_KEY"),
                    "GROK_API_KEY": os.getenv("GROK_API_KEY"),
                    "COHERE_API_KEY": os.getenv("COHERE_API_KEY"),
                    "GOOGLE_API_KEY": os.getenv("GOOGLE_API_KEY"),
                    "GOOGLE_SEARCH_ENGINE_ID": os.getenv("GOOGLE_SEARCH_ENGINE_ID")
                }

            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)

            # Validate config
            required_fields = [
                "OPENAI_API_KEY",
                "GEMINI_API_KEY",
                "GROQ_API_KEY",
                "GROK_API_KEY",
                "COHERE_API_KEY",
                "GOOGLE_API_KEY",
                "GOOGLE_SEARCH_ENGINE_ID"
            ]

            missing_fields = [field for field in required_fields if field not in config]
            if missing_fields:
                logger.warning(f"Missing fields in config: {missing_fields}")
                for field in missing_fields:
                    env_value = os.getenv(field)
                    if env_value:
                        config[field] = env_value
                        logger.info(f"Retrieved {field} from environment variables")

            return config
        except Exception as e:
            logger.error(f"Error loading config: {e}")
            raise

    def _extract_tables_from_pdf(self, pdf_path: str) -> List[Tuple[pd.DataFrame, int, Dict]]:
        """Extract tables from PDF while preserving structure."""
        tables = []

        try:
            # Try PyMuPDF first for better table extraction
            try:
                doc = fitz.open(pdf_path)
                for page_num, page in enumerate(doc):
                    try:
                        tables_on_page = page.find_tables()
                        for table_idx, table in enumerate(tables_on_page):
                            try:
                                table_dict = table.extract()
                                if table_dict and len(table_dict) > 1:  # Skip empty tables
                                    headers = table_dict[0]
                                    data = table_dict[1:]

                                    # Create DataFrame
                                    df = pd.DataFrame(data, columns=headers)

                                    # Store table metadata - with safer attribute access
                                    table_meta = {
                                        "page": page_num,
                                        "table_index": table_idx
                                    }

                                    # Safely add rect if available
                                    if hasattr(table, 'rect') and hasattr(table.rect, 'irect'):
                                        table_meta["rect"] = table.rect.irect

                                    tables.append((df, page_num, table_meta))
                            except Exception as table_err:
                                logger.warning(f"Error processing table {table_idx} on page {page_num}: {table_err}")
                    except Exception as page_err:
                        logger.warning(f"Error finding tables on page {page_num}: {page_err}")

                doc.close()

                if tables:
                    logger.info(f"Extracted {len(tables)} tables using PyMuPDF")
                    return tables
            except Exception as pymupdf_err:
                logger.warning(f"PyMuPDF table extraction failed: {pymupdf_err}, falling back to pdfplumber")

            # Fallback to pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table and len(table) > 1:  # Skip empty tables
                            headers = table[0]

                            # Clean up headers
                            headers = [str(h).strip() if h is not None else f"Column_{i}"
                                      for i, h in enumerate(headers)]

                            # Handle missing data
                            data = []
                            for row in table[1:]:
                                cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                data.append(cleaned_row)

                            # Create DataFrame
                            df = pd.DataFrame(data, columns=headers)

                            # Store table metadata
                            table_meta = {
                                "page": page_num,
                                "table_index": table_idx
                            }

                            tables.append((df, page_num, table_meta))

            if tables:
                logger.info(f"Extracted {len(tables)} tables using pdfplumber")
        except Exception as e:
            logger.error(f"Error extracting tables from PDF: {e}")

        return tables

    def _process_table(self, df: pd.DataFrame) -> str:
        """Convert table to a structured text format with improved formatting."""
        try:
            # Handle case where DataFrame columns are not unique
            if df.columns.duplicated().any():
                # Make columns unique by adding suffixes
                new_cols = []
                seen = {}
                for col in df.columns:
                    col_str = str(col)
                    if col_str in seen:
                        seen[col_str] += 1
                        new_cols.append(f"{col_str}_{seen[col_str]}")
                    else:
                        seen[col_str] = 0
                        new_cols.append(col_str)
                df.columns = new_cols
            else:
                # Ensure all column names are strings
                df.columns = df.columns.astype(str)

            # Drop completely empty columns
            df = df.dropna(axis=1, how='all')

            # Replace NaN with empty string
            df = df.fillna("")

            # Build table text representation
            text_parts = []

            # Add table header
            text_parts.append("TABLE DATA:")
            text_parts.append("Columns: " + ", ".join(df.columns))

            # Add row data with column names
            for idx, row in df.iterrows():
                row_text = []
                for col in df.columns:
                    try:
                        value = str(row[col]).strip()
                        if value and value.lower() not in ['nan', 'none', '']:
                            row_text.append(f"{col}: {value}")
                    except Exception as e:
                        logger.warning(f"Error processing column {col}: {e}")

                if row_text:  # Only add non-empty rows
                    text_parts.append("ROW " + str(idx+1) + ": " + " | ".join(row_text))

            # Handle special case of empty table
            if len(text_parts) <= 2:
                text_parts.append("(Empty table)")

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error processing table: {e}")
            return "Error processing table"
        except Exception as e:
            logger.error(f"Error processing table: {e}")
            return ""

    def add_file(self, file_path: str, track_versions: bool = True, progress_callback=None) -> bool:
        """Add or update a file in the knowledge base with improved processing and version tracking."""
        try:
            file_path = Path(file_path)
            file_name = file_path.name

            if progress_callback:
                progress_callback(f"Processing {file_name}...", 0)

            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                if progress_callback:
                    progress_callback(f"File not found: {file_name}", 100)
                return False

            if progress_callback:
                progress_callback(f"Calculating hash for {file_name}...", 10)

            # Calculate file hash
            file_hash = self._calculate_file_hash(file_path)

            # Track version history if enabled
            previous_versions = []
            if track_versions and file_name in self.file_registry:
                old_hash = self.file_registry[file_name].get('hash')
                if old_hash:
                    # Check if this document has previous versions tracked
                    if 'previous_versions' in self.file_registry[file_name]:
                        previous_versions = self.file_registry[file_name]['previous_versions']

                    # Add current version to history if file is being updated
                    if old_hash != file_hash:
                        previous_versions.append(old_hash)
                        # Keep only last 5 versions to avoid excessive growth
                        previous_versions = previous_versions[-5:]

            # Check if file has changed
            if file_name in self.file_registry and self.file_registry[file_name]['hash'] == file_hash:
                logger.info(f"File {file_name} already indexed and unchanged.")
                if progress_callback:
                    progress_callback(f"{file_name} already indexed (unchanged)", 100)
                return True

            if progress_callback:
                progress_callback(f"Extracting content from {file_name}...", 20)

            timestamp = int(time.time() * 1000)
            chunks = []
            metadatas = []

            file_extension = file_path.suffix.lower()

            # Process based on file type
            if file_extension == '.pdf':
                self._process_pdf_file(file_path, file_name, timestamp, chunks, metadatas, progress_callback)
            elif file_extension == '.docx':
                self._process_docx_file(file_path, file_name, timestamp, chunks, metadatas, progress_callback)
            elif file_extension in ['.xlsx', '.xls']:
                self._process_excel_file(file_path, file_name, timestamp, chunks, metadatas, progress_callback)
            elif file_extension == '.csv':
                self._process_csv_file(file_path, file_name, timestamp, chunks, metadatas, progress_callback)
            elif file_extension in ['.html', '.htm']:
                self._process_html_file(file_path, file_name, timestamp, chunks, metadatas, progress_callback)
            elif file_extension == '.txt':
                self._process_text_file(file_path, file_name, timestamp, chunks, metadatas, progress_callback)
            else:
                logger.warning(f"Unsupported file extension {file_extension}")
                if progress_callback:
                    progress_callback(f"Unsupported file type: {file_extension}", 100)
                return False

            if not chunks:
                logger.warning(f"No content extracted from {file_name}")
                if progress_callback:
                    progress_callback(f"No content extracted from {file_name}", 100)
                return False

            if progress_callback:
                progress_callback(f"Generating embeddings for {len(chunks)} chunks from {file_name}...", 60)

            # Calculate file hash
            file_hash = self._calculate_file_hash(str(file_path))

            # Get embeddings for all chunks using safe method
            logger.info(f"Generating embeddings for {len(chunks)} chunks from {file_name}")
            try:
                if self.ultra_safe_mode:
                    logger.warning("ULTRA-SAFE MODE: Using zero embeddings for all documents to avoid any crashes")
                    embeddings = np.zeros((len(chunks), self.dimension))
                else:
                    logger.info(f"Generating embeddings for {len(chunks)} chunks from {file_name} using {self.embedding_model_name} on {self.embedding_device_setting}")
                    try:
                        # Use the internal _get_embeddings method now that all flags are managed internally
                        embeddings = self._get_embeddings(chunks)
                        logger.info(f"Successfully generated embeddings for {file_name}")
                    except Exception as e:
                        logger.error(f"Error generating embeddings for {file_name}: {e}")
                        logger.error(traceback.format_exc())
                        logger.warning(f"Using zero embeddings for {file_name} as fallback due to embedding error")
                        embeddings = np.zeros((len(chunks), self.dimension))
            except Exception as e:
                logger.error(f"Error generating embeddings for {file_name}: {e}")
                logger.error(traceback.format_exc())
                # Create zero embeddings as fallback
                logger.warning(f"Using zero embeddings for {file_name} as fallback")
                embeddings = np.zeros((len(chunks), self.dimension))

            if progress_callback:
                progress_callback(f"Adding {file_name} to knowledge base index...", 80)

            # Remove old chunks if file was previously indexed
            if file_name in self.file_registry:
                self._remove_file_chunks(file_name)

            # Add new chunks to index
            self.index.add(embeddings.astype(np.float32))
            self.chunks.extend(chunks)
            self.metadatas.extend(metadatas)

            if progress_callback:
                progress_callback(f"Saving knowledge base updates...", 90)

            # Update registry with version tracking
            file_stat = file_path.stat()
            self.file_registry[file_name] = {
                "hash": file_hash,
                "path": str(file_path),
                "chunk_count": len(chunks),
                "timestamp": timestamp,
                "last_modified": int(file_stat.st_mtime),
                "size": file_stat.st_size,
                "embedding_model": self.embedding_model_name,
                "provider": self.embedding_provider.value,
                "previous_versions": previous_versions,
                "version_count": len(previous_versions) + 1
            }

            # Save updated state
            self.save_index()
            self._save_file_registry()
            self._save_embeddings_cache()

            if progress_callback:
                progress_callback(f"Successfully indexed {file_name} with {len(chunks)} chunks", 100)

            logger.info(f"Successfully indexed {file_name} with {len(chunks)} chunks")
            return True

        except Exception as e:
            logger.error(f"Error adding file {file_path}: {str(e)}")
            logger.debug(traceback.format_exc())
            if progress_callback:
                progress_callback(f"Error processing {file_path.name}: {str(e)}", 100)
            return False

    def _process_pdf_file(self, file_path: Path, file_name: str, timestamp: int,
                         chunks: List[str], metadatas: List[ChunkMetadata], progress_callback=None) -> None:
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                if progress_callback:
                    progress_callback(f"Parsing {total_pages} pages", 10)

                full_text = ""
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n"
                    if progress_callback:
                        progress = 10 + int(((i + 1) / total_pages) * 80)
                        progress_callback(f"Parsed page {i + 1}/{total_pages}", progress)
                
                chunked_texts = self._chunk_text(full_text)
                for i, chunk in enumerate(chunked_texts):
                    metadatas.append(ChunkMetadata(file_name=file_name, page_num=0, timestamp=timestamp, source_type='pdf', chunk_index=i, is_table=False))
                chunks.extend(chunked_texts)
                
                if progress_callback:
                    progress_callback("Finished parsing PDF", 100)

        except Exception as e:
            logger.error(f"Error processing PDF {file_name}: {e}")
            if progress_callback:
                progress_callback(f"Error: {e}", 100)

    def _process_docx_file(self, file_path: Path, file_name: str, timestamp: int,
                          chunks: List[str], metadatas: List[ChunkMetadata], progress_callback=None) -> None:
        """Process DOCX file with structure preservation."""
        try:
            doc = docx.Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])

            # Process document properties
            properties = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
                "subject": doc.core_properties.subject,
                "category": doc.core_properties.category,
                "keywords": doc.core_properties.keywords,
            }

            # Process document structure with headings and styles
            current_section = None
            current_text_block = []
            section_blocks = []

            # Extract paragraphs with heading information
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a heading
                is_heading = para.style.name.startswith('Heading')

                if is_heading:
                    # Save previous section if it exists
                    if current_text_block:
                        section_blocks.append((current_section, '\n'.join(current_text_block)))
                        current_text_block = []

                    current_section = text
                else:
                    current_text_block.append(text)

            # Save the last section
            if current_text_block:
                section_blocks.append((current_section, '\n'.join(current_text_block)))

            # Process each section
            for section_title, section_text in section_blocks:
                section_chunks = self._chunk_text(section_text)

                for i, chunk in enumerate(section_chunks):
                    chunks.append(chunk)
                    metadatas.append(ChunkMetadata(
                        file_name=file_name,
                        page_num=0,  # DOCX doesn't have pages in the same way
                        chunk_index=i,
                        is_table=False,
                        timestamp=timestamp,
                        source_type='docx',
                        section_title=section_title,
                        keywords=self.text_processor.extract_keywords(chunk)
                    ))

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                headers = []

                # Get headers from first row
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]

                # Get data from remaining rows
                for row_idx, row in enumerate(table.rows[1:], 1):
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(cell for cell in row_data):  # Skip empty rows
                        table_data.append(row_data)

                if headers and table_data:
                    # Convert to DataFrame
                    df = pd.DataFrame(table_data, columns=headers)
                    table_text = self._process_table(df)
                    table_chunks = self._chunk_text(table_text)

                    for i, chunk in enumerate(table_chunks):
                        chunks.append(chunk)
                        metadatas.append(ChunkMetadata(
                            file_name=file_name,
                            page_num=0,
                            chunk_index=i,
                            is_table=True,
                            timestamp=timestamp,
                            source_type='docx',
                            table_data=df.to_dict(),
                            keywords=self.text_processor.extract_keywords(chunk)
                        ))

            logger.info(f"Processed DOCX file {file_name} with {len(section_blocks)} sections and {len(doc.tables)} tables")

        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            logger.error(traceback.format_exc())
            raise

    def _process_excel_file(self, file_path: Path, file_name: str, timestamp: int,
                           chunks: List[str], metadatas: List[ChunkMetadata], progress_callback=None) -> None:
        """Process Excel file (ALL sheets)."""
        try:
            # Use pandas to read all sheets
            xlsx = pd.ExcelFile(file_path)
            sheet_names = xlsx.sheet_names

            logger.info(f"Processing {len(sheet_names)} sheets from Excel file {file_name}")

            for sheet_name in sheet_names:
                try:
                    # Read sheet data
                    df = pd.read_excel(xlsx, sheet_name=sheet_name)

                    if df.empty:
                        logger.info(f"Sheet '{sheet_name}' is empty, skipping")
                        continue

                    # Process sheet as table
                    sheet_text = f"SHEET: {sheet_name}\n" + self._process_table(df)
                    sheet_chunks = self._chunk_text(sheet_text)

                    # Add each chunk
                    for i, chunk in enumerate(sheet_chunks):
                        chunks.append(chunk)
                        metadatas.append(ChunkMetadata(
                            file_name=file_name,
                            page_num=0,
                            chunk_index=i,
                            is_table=True,
                            timestamp=timestamp,
                            source_type='excel',
                            sheet_name=sheet_name,
                            table_data=df.to_dict(),
                            keywords=self.text_processor.extract_keywords(chunk)
                        ))

                    # Additionally, process each row as individual chunk for better retrieval
                    row_chunks = self._process_excel_rows(df, sheet_name)

                    for i, (row_idx, row_text) in enumerate(row_chunks):
                        chunks.append(row_text)
                        row_data = df.iloc[row_idx].to_dict()
                        metadatas.append(ChunkMetadata(
                            file_name=file_name,
                            page_num=0,
                            chunk_index=i + len(sheet_chunks),  # Continue indexing
                            is_table=True,
                            timestamp=timestamp,
                            source_type='excel',
                            sheet_name=sheet_name,
                            row_data=row_data,
                            keywords=self.text_processor.extract_keywords(row_text)
                        ))

                except Exception as sheet_err:
                    logger.warning(f"Error processing sheet '{sheet_name}': {sheet_err}")

            # Try to extract additional metadata using openpyxl
            try:
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                properties = wb.properties.__dict__

                # Create metadata chunk
                meta_text = "Excel Document Properties:\n"
                for key, value in properties.items():
                    if value and not key.startswith('_'):
                        meta_text += f"{key}: {value}\n"

                if len(meta_text) > 50:  # Only add if we have meaningful metadata
                    chunks.append(meta_text)
                    metadatas.append(ChunkMetadata(
                        file_name=file_name,
                        page_num=0,
                        chunk_index=len(chunks) - 1,
                        is_table=False,
                        timestamp=timestamp,
                        source_type='excel',
                        section_title="Document Properties",
                        keywords=self.text_processor.extract_keywords(meta_text)
                    ))
            except Exception as meta_err:
                logger.warning(f"Could not extract Excel metadata: {meta_err}")

        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {e}")
            logger.error(traceback.format_exc())
            raise

    def _process_excel_rows(self, df: pd.DataFrame, sheet_name: str) -> List[Tuple[int, str]]:
        """Process individual rows from Excel for better retrieval."""
        row_chunks = []

        # Process rows individually for better recall during specific row searches
        for idx, row in df.iterrows():
            row_text = f"SHEET: {sheet_name}, ROW: {idx+1}\n"

            # Add column values with headers
            for col in df.columns:
                value = str(row[col]).strip()
                if value and value.lower() not in ['nan', 'none', '']:
                    row_text += f"{col}: {value}\n"

            if len(row_text) > 50:  # Only add non-empty rows
                row_chunks.append((idx, row_text))

        return row_chunks

    def _process_csv_file(self, file_path: Path, file_name: str, timestamp: int,
                         chunks: List[str], metadatas: List[ChunkMetadata], progress_callback=None) -> None:
        """Process CSV file."""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                result = chardet.detect(f.read(10000))
                encoding = result['encoding'] or 'utf-8'

            # Read CSV
            df = pd.read_csv(file_path, encoding=encoding, on_bad_lines='skip')

            if df.empty:
                logger.warning(f"CSV file {file_name} is empty")
                return

            # Process as table
            table_text = self._process_table(df)
            table_chunks = self._chunk_text(table_text)

            for i, chunk in enumerate(table_chunks):
                chunks.append(chunk)
                metadatas.append(ChunkMetadata(
                    file_name=file_name,
                    page_num=0,
                    chunk_index=i,
                    is_table=True,
                    timestamp=timestamp,
                    source_type='csv',
                    table_data=df.to_dict(),
                    keywords=self.text_processor.extract_keywords(chunk)
                ))

            # Process individual rows like in Excel
            row_chunks = self._process_excel_rows(df, "CSV")

            for i, (row_idx, row_text) in enumerate(row_chunks):
                chunks.append(row_text)
                row_data = df.iloc[row_idx].to_dict()
                metadatas.append(ChunkMetadata(
                    file_name=file_name,
                    page_num=0,
                    chunk_index=i + len(table_chunks),
                    is_table=True,
                    timestamp=timestamp,
                    source_type='csv',
                    row_data=row_data,
                    keywords=self.text_processor.extract_keywords(row_text)
                ))

        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            logger.error(traceback.format_exc())
            raise

    def _process_html_file(self, file_path: Path, file_name: str, timestamp: int,
                          chunks: List[str], metadatas: List[ChunkMetadata], progress_callback=None) -> None:
        """Process HTML file with structure preservation."""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                result = chardet.detect(f.read(10000))
                encoding = result['encoding'] or 'utf-8'

            # Read HTML
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                html_content = f.read()

            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')

            # Extract title
            title = ""
            if soup.title:
                title = soup.title.string

            # Extract metadata
            meta_tags = {}
            for meta in soup.find_all('meta'):
                if meta.get('name') and meta.get('content'):
                    meta_tags[meta['name']] = meta['content']

            # Extract headings
            headings = []
            for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                heading_text = h.get_text().strip()
                if heading_text:
                    headings.append((h.name, heading_text))

            # Extract tables
            tables = []
            for table in soup.find_all('table'):
                tables.append(table)

            # Process tables
            for i, table in enumerate(tables):
                try:
                    # Convert HTML table to DataFrame
                    table_rows = []
                    for tr in table.find_all('tr'):
                        row = []
                        for td in tr.find_all(['td', 'th']):
                            row.append(td.get_text().strip())
                        if row:
                            table_rows.append(row)

                    if table_rows:
                        if len(table_rows) > 1:
                            headers = table_rows[0]
                            data = table_rows[1:]
                            df = pd.DataFrame(data, columns=headers)

                            table_text = f"HTML TABLE {i+1}:\n" + self._process_table(df)
                            table_chunks = self._chunk_text(table_text)

                            for j, chunk in enumerate(table_chunks):
                                chunks.append(chunk)
                                metadatas.append(ChunkMetadata(
                                    file_name=file_name,
                                    page_num=0,
                                    chunk_index=len(chunks) - 1,
                                    is_table=True,
                                    timestamp=timestamp,
                                    source_type='html',
                                    table_data=df.to_dict(),
                                    keywords=self.text_processor.extract_keywords(chunk)
                                ))
                except Exception as table_err:
                    logger.warning(f"Error processing HTML table {i}: {table_err}")

            # Process main content
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()

            # Get text
            main_text = self.text_processor.html_to_text(str(soup))
            main_text = self.text_processor.clean_text(main_text)

            # Segment by headings where possible
            sections = []
            current_heading = title or "Main Content"
            current_content = []

            for line in main_text.split('\n'):
                # Check if line looks like a heading (e.g., # Heading or ## Heading from markdown)
                if re.match(r'^#{1,6}\s+', line):
                    # Save previous section
                    if current_content:
                        sections.append((current_heading, '\n'.join(current_content)))
                        current_content = []

                    # Update heading
                    current_heading = re.sub(r'^#{1,6}\s+', '', line)
                else:
                    current_content.append(line)

            # Save the last section
            if current_content:
                sections.append((current_heading, '\n'.join(current_content)))

            # Process each section
            for heading, content in sections:
                if not content.strip():
                    continue

                section_chunks = self._chunk_text(content)

                for i, chunk in enumerate(section_chunks):
                    chunks.append(chunk)
                    metadatas.append(ChunkMetadata(
                        file_name=file_name,
                        page_num=0,
                        chunk_index=len(chunks) - 1,
                        is_table=False,
                        timestamp=timestamp,
                        source_type='html',
                        section_title=heading,
                        keywords=self.text_processor.extract_keywords(chunk)
                    ))

        except Exception as e:
            logger.error(f"Error processing HTML file {file_path}: {e}")
            logger.error(traceback.format_exc())
            raise

    def _process_text_file(self, file_path: Path, file_name: str, timestamp: int,
                          chunks: List[str], metadatas: List[ChunkMetadata], progress_callback=None) -> None:
        """Process plain text file."""
        try:
            if progress_callback:
                progress_callback(f"Reading text file {file_name}...", 30)
                
            # Read with encoding detection
            text = self._read_text_with_encoding(file_path)

            if not text.strip():
                logger.warning(f"Text file {file_name} is empty")
                return

            if progress_callback:
                progress_callback(f"Processing content from {file_name}...", 40)
                
            # Clean text
            text = self.text_processor.clean_text(text)

            # Detect section titles
            section_titles = self.text_processor.detect_section_titles(text)

            # Chunk with section awareness
            if section_titles:
                chunks_with_sections = self._chunk_text_with_sections(text, section_titles)

                for i, (chunk_text, section) in enumerate(chunks_with_sections):
                    chunks.append(chunk_text)
                    metadatas.append(ChunkMetadata(
                        file_name=file_name,
                        page_num=0,
                        chunk_index=i,
                        is_table=False,
                        timestamp=timestamp,
                        source_type='text',
                        section_title=section,
                        keywords=self.text_processor.extract_keywords(chunk_text)
                    ))
            else:
                # Basic chunking
                text_chunks = self._chunk_text(text)

                for i, chunk in enumerate(text_chunks):
                    chunks.append(chunk)
                    metadatas.append(ChunkMetadata(
                        file_name=file_name,
                        page_num=0,
                        chunk_index=i,
                        is_table=False,
                        timestamp=timestamp,
                        source_type='text',
                        keywords=self.text_processor.extract_keywords(chunk)
                    ))

        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            logger.error(traceback.format_exc())
            raise

    def _read_text_with_encoding(self, file_path: Path) -> str:
        """Read text file with encoding detection."""
        try:
            # Try to detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB to detect encoding
                result = chardet.detect(raw_data)
                detected_encoding = result['encoding']

            if detected_encoding:
                with open(file_path, 'r', encoding=detected_encoding, errors='replace') as f:
                    return f.read()

            # Fallback encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            # Last resort
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()

        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    def _chunk_text_contextual(self, text: str) -> List[str]:
        """Enhanced contextual chunking that preserves context between chunks."""
        try:
            # Identify important sections and natural breaks
            section_titles = self.text_processor.detect_section_titles(text)
            paragraphs = re.split(r'\n\s*\n', text)

            # Prepare result chunks
            chunks = []

            # If very short text, return as is
            if len(text.split()) <= self.chunk_size:
                return [text]

            # If we have distinct paragraphs, use them as base units
            if len(paragraphs) > 1:
                current_chunk = []
                current_size = 0

                for para in paragraphs:
                    para_size = len(para.split())

                    # If paragraph alone exceeds chunk size, split it
                    if para_size > self.chunk_size:
                        # If we have content already, add it as a chunk
                        if current_chunk:
                            chunks.append("\n\n".join(current_chunk))
                            current_chunk = []
                            current_size = 0

                        # Split large paragraph using semantic chunking
                        para_chunks = self._chunk_text_semantic(para)
                        chunks.extend(para_chunks)
                    # Regular paragraph handling
                    elif current_size + para_size <= self.chunk_size:
                        current_chunk.append(para)
                        current_size += para_size
                    else:
                        # End current chunk and start new one
                        chunks.append("\n\n".join(current_chunk))
                        current_chunk = [para]
                        current_size = para_size

                # Add final chunk if any content remains
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
            else:
                # No clear paragraph structure, fall back to semantic chunking
                chunks = self._chunk_text_semantic(text)

            # Add overlap context to chunks (for metadata)
            result_chunks = []
            for chunk in chunks:
                result_chunks.append(chunk)

            # For very small results, fall back to semantic chunking
            if not result_chunks:
                return self._chunk_text_semantic(text)

            return result_chunks

        except Exception as e:
            logger.warning(f"Error in contextual chunking: {e}, falling back to semantic chunking")
            return self._chunk_text_semantic(text)

    def _chunk_text_with_sections(self, text: str, section_titles: List[Tuple[int, str]]) -> List[Tuple[str, str]]:
        """Enhanced section-based chunking with better context preservation."""
        try:
            # If no section titles, use regular chunking
            if not section_titles:
                return [(chunk, None) for chunk in self._chunk_text(text)]

            chunks_with_sections = []

            # Sort section titles by position
            section_titles.sort(key=lambda x: x[0])

            # Add text end as last boundary
            section_boundaries = section_titles + [(len(text), "")]

            # First pass - collect all sections
            sections = []
            for i in range(len(section_titles)):
                section_start, section_title = section_titles[i]
                section_end = section_boundaries[i+1][0]
                section_text = text[section_start:section_end]
                sections.append((section_title, section_text))

            # Second pass - process with awareness of surrounding sections
            for i, (section_title, section_text) in enumerate(sections):
                # Get context of adjacent sections for improved chunking
                context_before = sections[i-1][1][-100:] if i > 0 else ""
                context_after = sections[i+1][1][:100] if i < len(sections)-1 else ""

                # Use contextual chunking if available, otherwise use regular chunking
                try:
                    if self.chunking_strategy == ChunkingStrategy.CONTEXTUAL:
                        section_chunks = self._chunk_text_contextual(section_text)
                    else:
                        section_chunks = self._chunk_text(section_text)
                except Exception as e:
                    logger.warning(f"Error chunking section '{section_title}': {e}, using basic chunking")
                    section_chunks = self._chunk_text_fixed_size(section_text)

                # Add chunks with section title and context
                for chunk in section_chunks:
                    chunks_with_sections.append((chunk, section_title))

            return chunks_with_sections
        except Exception as e:
            logger.warning(f"Error in section-based chunking: {e}, falling back to basic chunking")
            return [(chunk, None) for chunk in self._chunk_text(text)]

    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text using selected strategy."""
        if not text or not text.strip():
            return []

        # Clean text
        text = self.text_processor.clean_text(text)

        if self.chunking_strategy == ChunkingStrategy.FIXED_SIZE:
            return self._chunk_text_fixed_size(text)
        elif self.chunking_strategy == ChunkingStrategy.SEMANTIC:
            return self._chunk_text_semantic(text)
        elif self.chunking_strategy == ChunkingStrategy.RECURSIVE:
            return self._chunk_text_recursive(text)
        elif self.chunking_strategy == ChunkingStrategy.SLIDING_WINDOW:
            return self._chunk_text_sliding_window(text)
        elif self.chunking_strategy == ChunkingStrategy.CONTEXTUAL:
            return self._chunk_text_contextual(text)
        else:
            # Default to semantic chunking
            return self._chunk_text_semantic(text)

    def _chunk_text_fixed_size(self, text: str) -> List[str]:
        """Simple fixed-size chunking with overlap."""
        sentences = tokenize_text(text)
        chunks = []
        current_chunk = []
        current_size = 0

        for sentence in sentences:
            # Count words for size estimation
            sentence_size = len(sentence.split())

            # If adding this sentence would exceed chunk size
            if current_size + sentence_size > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(" ".join(current_chunk))

                # Calculate overlap: take the last N tokens based on overlap ratio
                overlap_size = min(len(current_chunk),
                                  max(1, int(self.chunk_overlap / (current_size/len(current_chunk)))))

                # Start new chunk with overlap
                current_chunk = current_chunk[-overlap_size:]
                current_size = sum(len(s.split()) for s in current_chunk)

            # Add sentence to current chunk
            current_chunk.append(sentence)
            current_size += sentence_size

        # Add final chunk if not empty
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    def _chunk_text_semantic(self, text: str) -> List[str]:
        """Improved chunking with semantic boundaries."""
        sentences = tokenize_text(text)
        chunks = []
        current_chunk = []
        current_size = 0

        # Use sliding window with semantic boundaries
        for i, sentence in enumerate(sentences):
            sentence_size = len(sentence.split())

            # Check if adding this sentence would exceed chunk size
            if current_size + sentence_size > self.chunk_size:
                # If we have content, save the chunk
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                    # Start new chunk with overlap
                    overlap_start = max(0, len(current_chunk) - int(self.chunk_overlap / (current_size/len(current_chunk))))
                    current_chunk = current_chunk[overlap_start:]
                    current_size = sum(len(s.split()) for s in current_chunk)

            current_chunk.append(sentence)
            current_size += sentence_size

            # Check for better semantic boundaries
            if i < len(sentences) - 1:
                next_sentence = sentences[i+1]

                # Check for paragraph breaks or section headers
                if (re.search(r'\n{2,}', sentence) or
                    re.match(r'^\s*(?:[A-Z][A-Z\s]+|#+)\s*$', next_sentence) or
                    (current_size >= self.chunk_size * 0.8 and  # At least 80% full
                     re.search(r'[.!?]$', sentence))):  # Complete sentence ending

                    if current_chunk:
                        chunks.append(" ".join(current_chunk))
                        current_chunk = []
                        current_size = 0

        # Add final chunk if not empty
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    def _chunk_text_recursive(self, text: str) -> List[str]:
        """Recursive text chunking based on hierarchical boundaries."""
        # Define boundary types in order of importance
        boundary_patterns = [
            (r'\n\s*#{1,3}\s+[^\n]+\n', r'\n\s*#{1,3}\s+', 1.0),   # Markdown headings level 1-3
            (r'\n\s*#{4,6}\s+[^\n]+\n', r'\n\s*#{4,6}\s+', 0.8),   # Markdown headings level 4-6
            (r'\n\s*[A-Z][A-Z\s]+\n', r'\n\s*[A-Z][A-Z\s]+', 0.8), # ALL CAPS HEADINGS
            (r'\n\n+', r'\n\n+', 0.6),                             # Multiple newlines
            (r'(?<=[.!?])\s+(?=[A-Z])', r'(?<=[.!?])\s+', 0.4),    # Sentence boundaries
        ]

        def split_by_pattern(text_chunk, min_size, max_size, pattern, pattern_clean, importance):
            """Split text by a specific pattern."""
            # If chunk is small enough, don't split further
            if len(text_chunk.split()) <= max_size:
                return [text_chunk]

            # Try to split by pattern
            splits = re.split(pattern, text_chunk)

            # If split worked and created reasonable chunks
            if len(splits) > 1 and all(len(s.split()) >= min_size*importance for s in splits if s.strip()):
                # Clean up the splits (remove heading markers, etc)
                cleaned_splits = []
                for s in splits:
                    # Don't add empty splits
                    if not s.strip():
                        continue
                    # Clean up the split if needed
                    cleaned = re.sub(pattern_clean, '', s).strip()
                    if cleaned:
                        cleaned_splits.append(cleaned)
                return cleaned_splits
            else:
                return [text_chunk]  # No good split found

        def recursive_split(text_chunk, depth=0):
            """Recursively split text into chunks."""
            # Base case: chunk is small enough
            word_count = len(text_chunk.split())
            if word_count <= self.chunk_size or depth > 5:  # Prevent infinite recursion
                return [text_chunk]

            # Try each boundary pattern in order
            for pattern, pattern_clean, importance in boundary_patterns:
                chunks = split_by_pattern(
                    text_chunk,
                    min_size=self.chunk_size//4,  # Minimum quarter chunk
                    max_size=self.chunk_size,
                    pattern=pattern,
                    pattern_clean=pattern_clean,
                    importance=importance
                )

                # If we got multiple chunks, recurse on each one
                if len(chunks) > 1:
                    result = []
                    for chunk in chunks:
                        result.extend(recursive_split(chunk, depth+1))
                    return result

            # If no pattern worked well, use basic splitting
            return self._chunk_text_fixed_size(text_chunk)

        # Start recursive splitting
        return recursive_split(text)

    def _chunk_text_sliding_window(self, text: str) -> List[str]:
        """Sliding window chunking with variable stride."""
        tokens = text.split()
        chunks = []

        # Define stride (step size)
        stride = self.chunk_size - self.chunk_overlap
        stride = max(1, stride)  # Ensure stride is at least 1

        # Slide window and extract chunks
        for i in range(0, len(tokens), stride):
            # Get chunk tokens
            chunk_tokens = tokens[i:i + self.chunk_size]

            # Skip if too small
            if len(chunk_tokens) < self.chunk_size // 4 and i > 0:
                # If this is the last chunk and it's small, append to previous
                if chunks:
                    last_chunk = chunks[-1].split()
                    combined = last_chunk + chunk_tokens
                    # Replace last chunk with combined if not too large
                    if len(combined) <= self.chunk_size * 1.5:
                        chunks[-1] = " ".join(combined)
                continue

            # Add chunk
            chunks.append(" ".join(chunk_tokens))

        return chunks

    def get_relevant_chunks(self,
                           query: str,
                           n_results: int = 10,
                           alpha: float = 0.5,
                           filter_criteria: Optional[Dict[str, Any]] = None,
                           reranking: bool = True,
                           cross_encoder_reranking: bool = False,
                           query_expansion: bool = True) -> List[Dict[str, Any]]:
        """Ultra-safe version of hybrid search to avoid crashes."""
        try:
            logger.info(f"RAGHandler.get_relevant_chunks: Starting retrieval for query '{query[:50]}...' (n_results={n_results}, alpha={alpha})")
            
            # Apply safe retrieval mode modifications
            if self.safe_retrieval_mode:
                logger.info("RAGHandler: Safe retrieval mode is ON. Modifying search parameters.")
                # Example modifications for safe mode:
                n_results = max(1, n_results // 2)  # Reduce number of results
                alpha = 0.7  # Favor semantic search more, potentially less keyword noise
                # Disable query expansion if it's too aggressive for safe mode
                if query_expansion:
                    logger.info("RAGHandler: Disabling query expansion in safe retrieval mode.")
                    query_expansion = False 
                # Could also tighten filter_criteria, e.g., higher importance_score
                if filter_criteria and "importance_score" in filter_criteria:
                    filter_criteria["importance_score"] = min(1.0, filter_criteria["importance_score"] * 1.2) # Be stricter
                elif filter_criteria is None:
                    filter_criteria = {"importance_score": 0.6} # Default stricter filter
                else: # filter_criteria exists but no importance_score
                    filter_criteria["importance_score"] = 0.6
            
            if not self.chunks:
                logger.warning("RAGHandler.get_relevant_chunks: No chunks in index.")
                return []

            if not query.strip():
                logger.warning("RAGHandler.get_relevant_chunks: Empty query provided.")
                return []

            # Clean and preprocess query
            cleaned_query = self.text_processor.clean_text(query)
            logger.info(f"RAGHandler.get_relevant_chunks: Cleaned query: '{cleaned_query[:50]}...'")

            # Query expansion
            if query_expansion:
                expanded_queries = self._expand_query(cleaned_query)
                logger.info(f"RAGHandler.get_relevant_chunks: Query expansion generated {len(expanded_queries)} additional queries")
            else:
                expanded_queries = [cleaned_query]

            # Calculate k for semantic search (use more candidates for better results)
            k_semantic = min(n_results * 3, len(self.chunks))
            logger.info(f"RAGHandler.get_relevant_chunks: Using k_semantic={k_semantic} for initial search")

            # Step 1: Semantic search
            all_candidate_indices = set()
            all_scores = {}

            for expanded_query in expanded_queries:
                try:
                    # Get embeddings for the query
                    query_embedding = self._get_embeddings([expanded_query])[0]
                    
                    # Search the index
                    if self.index.ntotal > 0:
                        scores, indices = self.index.search(query_embedding.reshape(1, -1), k_semantic)
                        
                        # Add results to candidates
                        for score, idx in zip(scores[0], indices[0]):
                            if idx != -1:  # Valid index
                                all_candidate_indices.add(idx)
                                # Keep the best score for each index
                                if idx not in all_scores or score > all_scores[idx]:
                                    all_scores[idx] = float(score)
                        
                        logger.info(f"RAGHandler.get_relevant_chunks: Semantic search for '{expanded_query[:30]}...' found {len([i for i in indices[0] if i != -1])} candidates")
                    else:
                        logger.warning("RAGHandler.get_relevant_chunks: Index is empty")
                        
                except Exception as e:
                    logger.error(f"RAGHandler.get_relevant_chunks: Error in semantic search for query '{expanded_query[:30]}...': {str(e)}")
                    continue

            # Convert to sorted list of indices based on score
            candidate_indices = sorted(
                list(all_candidate_indices),
                key=lambda idx: all_scores.get(idx, 0),
                reverse=True
            )[:k_semantic]

            if not candidate_indices:
                logger.info(f"RAGHandler.get_relevant_chunks: No candidates found in initial semantic search for query '{query[:50]}...'")
                return []

            logger.info(f"RAGHandler.get_relevant_chunks: Initial semantic search yielded {len(candidate_indices)} unique candidates")

            # Step 2: Filter results
            if filter_criteria:
                filtered_indices = []
                for idx in candidate_indices:
                    if 0 <= idx < len(self.metadatas):
                        metadata = self.metadatas[idx]
                        
                        # Check importance score filter
                        if "importance_score" in filter_criteria:
                            min_importance = filter_criteria["importance_score"]
                            if metadata.importance_score < min_importance:
                                continue
                        
                        # Check language filter
                        if "language" in filter_criteria:
                            target_language = filter_criteria["language"]
                            if metadata.language and metadata.language != target_language:
                                continue
                        
                        filtered_indices.append(idx)
                
                candidate_indices = filtered_indices
                logger.info(f"RAGHandler.get_relevant_chunks: Filtering reduced candidates from {len(all_candidate_indices)} to {len(candidate_indices)}")

            if not candidate_indices:
                logger.info(f"RAGHandler.get_relevant_chunks: No candidates left after filtering for query '{query[:50]}...'")
                return []

            # Step 3: Hybrid reranking
            if reranking and len(candidate_indices) > 1:
                try:
                    # Prepare candidate texts for reranking
                    candidate_texts = []
                    valid_indices = []
                    
                    for idx in candidate_indices:
                        if 0 <= idx < len(self.chunks):
                            candidate_texts.append(self.chunks[idx])
                            valid_indices.append(idx)
                    
                    if candidate_texts:
                        # Calculate hybrid scores
                        hybrid_scores = []
                        
                        for i, idx in enumerate(valid_indices):
                            # Combine semantic score with text-based features
                            semantic_score = all_scores.get(idx, 0)
                            
                            # Simple text-based scoring (keyword matching, length, etc.)
                            text_score = 0
                            chunk_text = candidate_texts[i].lower()
                            query_terms = cleaned_query.lower().split()
                            
                            # Keyword matching
                            for term in query_terms:
                                if term in chunk_text:
                                    text_score += 1
                            
                            # Normalize text score
                            text_score = text_score / len(query_terms) if query_terms else 0
                            
                            # Combine scores with alpha parameter
                            hybrid_score = alpha * semantic_score + (1 - alpha) * text_score
                            hybrid_scores.append(hybrid_score)
                        
                        # Sort by hybrid scores
                        scored_pairs = list(zip(valid_indices, hybrid_scores))
                        scored_pairs.sort(key=lambda x: x[1], reverse=True)
                        
                        candidate_indices = [idx for idx, _ in scored_pairs]
                        logger.info(f"RAGHandler.get_relevant_chunks: Hybrid reranking completed for {len(candidate_indices)} candidates")
                        
                except Exception as e:
                    logger.error(f"RAGHandler.get_relevant_chunks: Error in hybrid reranking: {str(e)}")
                    # Fall back to semantic scores only
                    pass

            # Step 4: Cross-encoder reranking (if enabled and available)
            if cross_encoder_reranking and len(candidate_indices) > 1:
                try:
                    cross_encoder_scores = self._cross_encoder_rerank(cleaned_query, [self.chunks[idx] for idx in candidate_indices if 0 <= idx < len(self.chunks)])
                    if cross_encoder_scores is not None:
                        # Re-sort based on cross-encoder scores
                        scored_pairs = list(zip(candidate_indices, cross_encoder_scores))
                        scored_pairs.sort(key=lambda x: x[1], reverse=True)
                        candidate_indices = [idx for idx, _ in scored_pairs]
                        logger.info(f"RAGHandler.get_relevant_chunks: Cross-encoder reranking completed")
                except Exception as e:
                    logger.error(f"RAGHandler.get_relevant_chunks: Error in cross-encoder reranking: {str(e)}")

            # Take top n_results
            result_indices = candidate_indices[:n_results]
            result_scores = [all_scores.get(idx, 0) for idx in result_indices]

            # Log the final results
            if result_indices:
                logger.info(f"RAGHandler.get_relevant_chunks: Top {len(result_indices)} RAG retrieval results found for query '{query[:50]}...':")
                for i, idx in enumerate(result_indices):
                    # Add a check to ensure idx is within bounds
                    if 0 <= idx < len(self.metadatas):
                        file_name_info = self.metadatas[idx].file_name
                        score_info = result_scores[i]
                        chunk_preview = self.chunks[idx][:100].replace('\n', ' ') # Preview first 100 chars, strip newlines
                        logger.info(f"  Result {i+1}: Chunk ID {idx}, Score {score_info:.4f}, File: '{file_name_info}', Content Preview: '{chunk_preview}...'")
                    else:
                        logger.warning(f"RAGHandler.get_relevant_chunks: Invalid index {idx} in results (out of bounds).")
            else:
                logger.info(f"RAGHandler.get_relevant_chunks: RAG retrieval yielded no results for query '{query[:50]}...'")

            # Build results
            results = []
            for i, idx in enumerate(result_indices):
                if 0 <= idx < len(self.chunks) and 0 <= idx < len(self.metadatas):
                    result = {
                        'content': self.chunks[idx],
                        'metadata': self.metadatas[idx],
                        'score': result_scores[i]
                    }
                    results.append(result)

            logger.info(f"RAGHandler.get_relevant_chunks: Returning {len(results)} results for query '{query[:50]}...'")
            return results
            
        except Exception as e:
            logger.error(f"RAGHandler.get_relevant_chunks: Error querying knowledge base for '{query[:50]}...': {str(e)}")
            logger.error(traceback.format_exc())
            return []

    def save_index(self):
        """Save the FAISS index and associated data."""
        try:
            # Check if directory is writable before attempting to save
            if not self._is_directory_writable(self.persist_directory):
                logger.error(f"Cannot save index: Directory {self.persist_directory} is not writable")
                # Try to switch to a fallback directory
                home_dir = Path.home() / ".rag_knowledge_base"
                home_dir.mkdir(parents=True, exist_ok=True)

                # Create a unique fallback directory
                fallback_dir = home_dir / f"fallback_{uuid.uuid4().hex[:8]}"
                fallback_dir.mkdir(parents=True, exist_ok=True)

                logger.warning(f"Attempting to save to fallback directory: {fallback_dir}")

                # Update paths temporarily for this save operation
                old_persist_dir = self.persist_directory
                self.persist_directory = fallback_dir
                self.index_path = fallback_dir / "faiss.index"
                self.metadata_path = fallback_dir / "metadata.pkl"
                self.file_registry_path = fallback_dir / "file_registry.json"
                self.embeddings_cache_path = fallback_dir / "embeddings_cache.pkl"

                # Check if fallback is writable
                if not self._is_directory_writable(fallback_dir):
                    logger.error(f"Fallback directory {fallback_dir} is also not writable. Cannot save index.")
                    # Restore original paths
                    self.persist_directory = old_persist_dir
                    self.index_path = old_persist_dir / "faiss.index"
                    self.metadata_path = old_persist_dir / "metadata.pkl"
                    self.file_registry_path = old_persist_dir / "file_registry.json"
                    self.embeddings_cache_path = old_persist_dir / "embeddings_cache.pkl"
                    return

            # Save FAISS index
            faiss.write_index(self.index, str(self.index_path))

            # Save metadata and chunks
            with open(self.metadata_path, 'wb') as f:
                pickle.dump({
                    'chunks': self.chunks,
                    'metadatas': self.metadatas,
                    'dimension': self.dimension,
                    'embedding_model': self.embedding_model_name,
                    'embedding_provider': self.embedding_provider.value,
                    'version': '3.0'
                }, f)

            # Save embeddings cache separately
            self._save_embeddings_cache()

            logger.info(f"Index and metadata saved successfully to {self.persist_directory}")
        except (IOError, PermissionError) as e:
            logger.error(f"Permission error saving index and metadata: {e}")
            logger.error(traceback.format_exc())
            logger.warning("Your knowledge base was processed but could not be saved due to file permission issues.")
            logger.warning("Please select a writable directory for your knowledge base.")
        except Exception as e:
            logger.error(f"Error saving index and metadata: {e}")
            logger.error(traceback.format_exc())

    def _sync_file_registry_with_metadata(self) -> bool:
        """
        Synchronize the file registry with the actual metadata to prevent mismatches.
        This ensures that the file registry accurately reflects what's in the metadata.
        """
        try:
            if not self.metadatas:
                # No metadata, clear registry
                self.file_registry = {}
                self._save_file_registry()
                return True
            
            # Group metadata by file name
            from collections import defaultdict
            file_groups = defaultdict(list)
            for metadata in self.metadatas:
                file_groups[metadata.file_name].append(metadata)
            
            # Create new registry from metadata
            new_registry = {}
            current_time = int(time.time() * 1000)
            
            for filename, metadata_list in file_groups.items():
                # Get the first metadata entry for this file
                first_metadata = metadata_list[0]
                
                # Check if file still exists
                file_exists = os.path.exists(filename)
                
                # Create registry entry
                registry_entry = {
                    "hash": "unknown",  # We don't have the original hash
                    "path": filename,
                    "chunk_count": len(metadata_list),
                    "timestamp": current_time,
                    "last_modified": current_time,
                    "size": 0,  # We don't have the original size
                    "embedding_model": getattr(first_metadata, 'embedding_model', self.embedding_model_name),
                    "provider": getattr(first_metadata, 'provider', self.embedding_provider.value),
                    "previous_versions": [],
                    "version_count": 1,
                    "file_exists": file_exists
                }
                
                new_registry[filename] = registry_entry
            
            # Update the file registry
            self.file_registry = new_registry
            self._save_file_registry()
            
            logger.info(f"Synced file registry with metadata: {len(new_registry)} files")
            return True
            
        except Exception as e:
            logger.error(f"Error syncing file registry with metadata: {e}")
            return False

    def load_index(self):
        """Load the FAISS index and associated data if they exist."""
        try:
            if self.index_path.exists() and self.metadata_path.exists():
                # Check embedding model compatibility first
                if not self._check_embedding_model_compatibility():
                    logger.warning("Embedding model compatibility check failed.")
                    if not self._handle_embedding_model_change():
                        # If re-indexing is needed, return False to indicate empty state
                        return False
                
                # Check index compatibility
                try:
                    temp_index = faiss.read_index(str(self.index_path))
                    if temp_index.d != self.dimension:
                        logger.warning(f"Dimension mismatch: Index has {temp_index.d} dimensions, but current model uses {self.dimension}.")
                        # Handle dimension mismatch
                        if not self._handle_embedding_model_change():
                            return False
                        # After handling, create new empty index
                        self.index = self._create_index()
                        self.chunks = []
                        self.metadatas = []
                        return False
                except Exception as e:
                    logger.error(f"Error reading index: {e}")
                    # Initialize empty index
                    self.index = self._create_index()
                    self.chunks = []
                    self.metadatas = []
                    return False

                # Load FAISS index since it's compatible
                self.index = temp_index

                # Load metadata and chunks
                try:
                    with open(self.metadata_path, 'rb') as f:
                        data = pickle.load(f)
                        self.chunks = data['chunks']
                        self.metadatas = data['metadatas']

                        # Validate metadata
                        for i, metadata in enumerate(self.metadatas):
                            if not hasattr(metadata, 'source_type'):
                                # Try to infer source type from file extension
                                file_ext = os.path.splitext(metadata.file_name)[1].lower()
                                source_type_map = {
                                    '.pdf': 'pdf',
                                    '.docx': 'docx',
                                    '.xlsx': 'excel',
                                    '.xls': 'excel',
                                    '.csv': 'csv',
                                    '.html': 'html',
                                    '.htm': 'html',
                                    '.txt': 'text'
                                }
                                metadata.source_type = source_type_map.get(file_ext, 'unknown')

                        # Check for version compatibility
                        if 'version' in data and data['version'] == '3.0':
                            # New version metadata
                            loaded_dimension = data.get('dimension', self.dimension)
                            loaded_provider = data.get('embedding_provider')
                            loaded_model = data.get('embedding_model')

                            # Double-check dimension match
                            if loaded_dimension != self.dimension:
                                logger.warning(f"Metadata dimension ({loaded_dimension}) doesn't match index dimension ({self.index.d})")

                            if loaded_provider and loaded_provider != self.embedding_provider.value:
                                logger.warning(
                                    f"Loaded index was created with provider {loaded_provider}, "
                                    f"but current provider is {self.embedding_provider.value}"
                                )

                            if loaded_model and loaded_model != self.embedding_model_name:
                                logger.warning(
                                    f"Loaded index was created with model {loaded_model}, "
                                    f"but current model is {self.embedding_model_name}"
                                )
                except Exception as metadata_err:
                    logger.error(f"Error loading metadata: {metadata_err}")
                    # If we can't load metadata but have a valid index, we need to reset
                    self.index = self._create_index()
                    self.chunks = []
                    self.metadatas = []
                    return False

                # Load file registry if it exists
                self.file_registry = self._load_file_registry()
                
                # Sync file registry with metadata to prevent mismatches
                self._sync_file_registry_with_metadata()

                logger.info(f"Successfully loaded knowledge base with {len(self.chunks)} chunks from {self.persist_directory}")
                return True
            else:
                # Initialize empty index
                self.index = self._create_index()
                self.chunks = []
                self.metadatas = []
                self.file_registry = {}

                # Save empty state
                self.save_index()
                logger.info(f"Initialized new knowledge base at {self.persist_directory}")
                return False
        except Exception as e:
            logger.error(f"Error loading index and metadata from {self.persist_directory}: {e}")
            logger.error(traceback.format_exc())
            # Reinitialize empty index
            self.index = self._create_index()
            self.chunks = []
            self.metadatas = []
            self.file_registry = {}
            return False

    def _load_file_registry(self) -> Dict[str, Any]:
        """Load the file registry from disk or create a new one if it doesn't exist."""
        try:
            if os.path.exists(self.file_registry_path):
                with open(self.file_registry_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            else:
                logger.info("No existing file registry found. Creating new one.")
                return {}
        except Exception as e:
            logger.error(f"Error loading file registry: {e}")
            return {}

    def _save_file_registry(self) -> None:
        """Save the file registry to disk."""
        try:
            # Check if parent directory is writable
            if not self._is_directory_writable(self.file_registry_path.parent):
                logger.error(f"Cannot save file registry: Directory {self.file_registry_path.parent} is not writable")
                return

            with open(self.file_registry_path, 'w', encoding='utf-8') as f:
                json.dump(self.file_registry, f, indent=4)
        except (IOError, PermissionError) as e:
            logger.error(f"Permission error saving file registry: {e}")
        except Exception as e:
            logger.error(f"Error saving file registry: {e}")

    def check_and_repair_file_registry(self) -> bool:
        """Check file registry for embedding model consistency and repair if needed."""
        try:
            if not self.file_registry_path.exists():
                return True

            with open(self.file_registry_path, 'r', encoding='utf-8') as f:
                registry = json.load(f)

            needs_update = False
            for file_name, file_info in registry.items():
                # Check for model name mismatch
                if file_info.get('embedding_model', '') != self.embedding_model_name:
                    file_info['embedding_model'] = self.embedding_model_name
                    file_info['provider'] = self.embedding_provider.value
                    needs_update = True

                # Check if dimension is stored and needs update
                if 'dimension' in file_info and file_info['dimension'] != self.dimension:
                    file_info['dimension'] = self.dimension
                    needs_update = True

            if needs_update:
                logger.info(f"Updating file registry to use model {self.embedding_model_name} with dimension {self.dimension}")
                with open(self.file_registry_path, 'w', encoding='utf-8') as f:
                    json.dump(registry, f, indent=4)
                self.file_registry = registry

            return True
        except Exception as e:
            logger.error(f"Error checking/repairing file registry: {e}")
            return False

    def _load_embeddings_cache(self) -> Dict[str, np.ndarray]:
        """Load embeddings cache from disk."""
        try:
            if os.path.exists(self.embeddings_cache_path):
                with open(self.embeddings_cache_path, 'rb') as f:
                    cache = pickle.load(f)
                logger.info(f"Loaded embeddings cache with {len(cache)} entries")
                return cache
            else:
                return {}
        except Exception as e:
            logger.warning(f"Error loading embeddings cache: {e}")
            return {}

    def _save_embeddings_cache(self) -> None:
        """Save embeddings cache to disk."""
        try:
            # Check if parent directory is writable
            if not self._is_directory_writable(self.embeddings_cache_path.parent):
                logger.warning(f"Cannot save embeddings cache: Directory {self.embeddings_cache_path.parent} is not writable")
                return

            # Limit cache size to prevent excessive growth
            if len(self.embeddings_cache) > 100000:
                # Keep only the newest 80% of entries
                keep_count = int(len(self.embeddings_cache) * 0.8)
                self.embeddings_cache = dict(list(self.embeddings_cache.items())[-keep_count:])

            with open(self.embeddings_cache_path, 'wb') as f:
                pickle.dump(self.embeddings_cache, f)
        except (IOError, PermissionError) as e:
            logger.warning(f"Permission error saving embeddings cache: {e}")
        except Exception as e:
            logger.warning(f"Error saving embeddings cache: {e}")

    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of a file."""
        try:
            sha256_hash = hashlib.sha256()
            with open(file_path, "rb") as f:
                for byte_block in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(byte_block)
            return sha256_hash.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating file hash for {file_path}: {e}")
            raise

    def remove_file(self, file_name: str) -> bool:
        """Remove a file from the knowledge base."""
        try:
            if file_name in self.file_registry:
                self._remove_file_chunks(file_name)
                del self.file_registry[file_name]
                self._save_file_registry()
                self.save_index()
                logger.info(f"Successfully removed {file_name}")
                return True
            else:
                logger.warning(f"File {file_name} not found in the registry.")
                return False
        except Exception as e:
            logger.error(f"Error removing file {file_name}: {str(e)}")
            return False

    def _remove_file_chunks(self, file_name: str) -> None:
        """Remove chunks associated with a file from the index."""
        try:
            # Find indices of chunks to remove
            indices_to_remove = [
                i for i, metadata in enumerate(self.metadatas)
                if metadata.file_name == file_name
            ]

            if not indices_to_remove:
                return

            # Create a mask for chunks to keep
            keep_mask = np.ones(len(self.chunks), dtype=bool)
            keep_mask[indices_to_remove] = False

            # Create new index
            new_index = self._create_index()

            # Get embeddings for remaining chunks
            remaining_chunks = [self.chunks[i] for i in range(len(self.chunks)) if keep_mask[i]]
            if remaining_chunks:
                logger.info(f"Regenerating index without file {file_name}")
                embeddings = self._get_embeddings(remaining_chunks)
                new_index.add(embeddings.astype(np.float32))

            # Update instance variables
            self.index = new_index
            self.chunks = remaining_chunks
            self.metadatas = [m for i, m in enumerate(self.metadatas) if keep_mask[i]]

        except Exception as e:
            logger.error(f"Error removing chunks for {file_name}: {e}")
            raise

    def get_indexed_files(self) -> List[str]:
        """Get list of currently indexed files (backward compatible version)."""
        return list(self.file_registry.keys())

    def get_indexed_files_detailed(self) -> List[Dict[str, Any]]:
        """Get detailed information about currently indexed files."""
        result = []
        for file_name, info in self.file_registry.items():
            # Count chunks for this file
            chunk_count = sum(1 for metadata in self.metadatas if metadata.file_name == file_name)

            # Get file details
            file_info = {
                "file_name": file_name,
                "path": info.get("path", ""),
                "chunk_count": chunk_count,
                "timestamp": info.get("timestamp"),
                "last_modified": info.get("last_modified"),
                "size": info.get("size", 0),
                "embedding_model": info.get("embedding_model", self.embedding_model_name),
                "provider": info.get("provider", self.embedding_provider.value)
            }
            result.append(file_info)

        return result

    def clear_knowledge_base(self) -> bool:
        """Clear all documents from the knowledge base."""
        try:
            # Reset index
            self.index = self._create_index()
            self.chunks = []
            self.metadatas = []
            self.file_registry = {}

            # Save empty state
            self.save_index()
            self._save_file_registry()

            # Clear embeddings cache
            self.embeddings_cache = {}
            if self.embeddings_cache_path.exists():
                self.embeddings_cache_path.unlink()

            logger.info("Knowledge base cleared successfully.")
            return True
        except Exception as e:
            logger.error(f"Error clearing knowledge base: {str(e)}")
            return False

    def get_knowledge_base_path(self) -> str:
        """Returns the current persist directory path."""
        return str(self.persist_directory)

    def set_knowledge_base_path(self, path: str) -> bool:
        """Sets a new persist directory path and loads the knowledge base from that location."""
        try:
            # Save current state before switching if we have data
            if self.chunks and self.index.ntotal > 0:
                self.save_index()

            # Set up new directory
            new_path = Path(path)
            new_path.mkdir(parents=True, exist_ok=True)

            # Check if the directory is writable
            if not self._is_directory_writable(new_path):
                logger.warning(f"Directory {path} is not writable. Using local fallback directory.")
                # Create a fallback directory in the user's home directory
                home_dir = Path.home() / ".rag_knowledge_base"
                home_dir.mkdir(parents=True, exist_ok=True)

                # Create a subdirectory based on the original path name to avoid conflicts
                safe_name = new_path.name or "default"
                fallback_path = home_dir / safe_name
                fallback_path.mkdir(parents=True, exist_ok=True)

                logger.info(f"Using fallback directory: {fallback_path}")
                new_path = fallback_path

                # Verify the fallback is writable
                if not self._is_directory_writable(new_path):
                    logger.error(f"Fallback directory {fallback_path} is also not writable. Cannot proceed.")
                    return False

            # Update paths
            self.persist_directory = new_path
            self.index_path = new_path / "faiss.index"
            self.metadata_path = new_path / "metadata.pkl"
            self.file_registry_path = new_path / "file_registry.json"
            self.embeddings_cache_path = new_path / "embeddings_cache.pkl"

            # Reset embeddings cache
            self.embeddings_cache = self._load_embeddings_cache()

            # Check if index exists and is compatible with current model dimensions
            index_compatible = True
            if self.index_path.exists():
                try:
                    temp_index = faiss.read_index(str(self.index_path))
                    if temp_index.d != self.dimension:
                        logger.warning(f"Dimension mismatch: Index has {temp_index.d} dimensions, but current model uses {self.dimension}.")
                        if self._check_model_mismatch():
                            index_compatible = False
                except Exception as ex:
                    logger.error(f"Error checking index compatibility: {ex}")
                    index_compatible = False

            # If we have an incompatible index, we need to recreate it
            if not index_compatible and self.index_path.exists():
                logger.warning("Rebuilding index due to dimension or model mismatch...")
                # Backup old index and metadata files
                backup_suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
                try:
                    if self.index_path.exists():
                        self.index_path.rename(self.index_path.with_suffix(f".index.bak_{backup_suffix}"))
                    if self.metadata_path.exists():
                        self.metadata_path.rename(self.metadata_path.with_suffix(f".pkl.bak_{backup_suffix}"))
                    # Also create new index with correct dimensions
                    self.index = self._create_index()
                    self.chunks = []
                    self.metadatas = []
                    # Update embedding model in file registry
                    self._update_file_registry_model()
                    logger.info(f"Created new index with dimension {self.dimension}")
                except Exception as bex:
                    logger.error(f"Failed to backup/recreate index: {bex}")
                    return False

            # Try to load existing data
            if self.load_index():
                logger.info(f"Successfully loaded existing knowledge base from {path}")
            else:
                logger.info(f"Created new knowledge base at {path}")

            return True

        except Exception as e:
            logger.error(f"Error setting knowledge base path: {str(e)}")
            return False

    def _check_model_mismatch(self) -> bool:
        """Check if the file registry model differs from current model."""
        try:
            if self.file_registry_path.exists():
                with open(self.file_registry_path, 'r', encoding='utf-8') as f:
                    registry = json.load(f)

                # Check if any file has a different embedding model
                for file_info in registry.values():
                    if file_info.get('embedding_model', '') != self.embedding_model_name:
                        return True
            return False
        except Exception as e:
            logger.error(f"Error checking model mismatch: {e}")
            return False

    def _is_directory_writable(self, directory_path: Path) -> bool:
        """Check if a directory is writable by attempting to create a temporary file."""
        try:
            # Create a unique temporary filename
            test_file = directory_path / f"write_test_{uuid.uuid4().hex}.tmp"

            # Try to write to the file
            with open(test_file, 'w') as f:
                f.write('test')

            # Clean up
            test_file.unlink()
            return True
        except (IOError, PermissionError):
            return False
        except Exception as e:
            logger.error(f"Error checking directory permissions: {e}")
            return False

    def _update_file_registry_model(self):
        """Update embedding model info in file registry."""
        try:
            if not self.file_registry_path.exists():
                return

            with open(self.file_registry_path, 'r', encoding='utf-8') as f:
                registry = json.load(f)

            # Update embedding model for all files
            for file_name in registry:
                registry[file_name]['embedding_model'] = self.embedding_model_name
                registry[file_name]['provider'] = self.embedding_provider.value

            # Save updated registry
            with open(self.file_registry_path, 'w', encoding='utf-8') as f:
                json.dump(registry, f, indent=4)

            logger.info(f"Updated file registry to use model {self.embedding_model_name}")
        except Exception as e:
            logger.error(f"Error updating file registry: {e}")

    def batch_add_files(self, file_paths: List[str], max_workers: Optional[int] = None, progress_callback=None) -> Dict[str, bool]:
        """Add multiple files in parallel."""
        if not max_workers:
            max_workers = self.max_threads

        results = {}
        completed_count = 0
        total_files = len(file_paths)

        def file_progress_callback(message, percentage):
            """Individual file progress callback that updates overall progress"""
            if progress_callback:
                # Calculate overall progress based on completed files and current file progress
                overall_progress = (completed_count * 100 + percentage) / total_files
                progress_callback(message, int(overall_progress))

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {executor.submit(self.add_file, file_path, True, file_progress_callback): file_path for file_path in file_paths}
            
            if progress_callback:
                progress_callback(f"Starting processing of {total_files} files...", 0)
            
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    result = future.result()
                    results[os.path.basename(file_path)] = result
                    completed_count += 1
                    
                    if progress_callback:
                        progress_callback(f"Completed {completed_count}/{total_files} files", int((completed_count * 100) / total_files))
                        
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    results[os.path.basename(file_path)] = False
                    completed_count += 1
                    
                    if progress_callback:
                        progress_callback(f"Error processing {os.path.basename(file_path)}: {str(e)}", int((completed_count * 100) / total_files))

        if progress_callback:
            success_count = sum(1 for result in results.values() if result)
            progress_callback(f"Completed: {success_count}/{total_files} files successfully processed", 100)

        return results

    def add_directory(self, directory_path: str, recursive: bool = True,
                     file_extensions: Optional[List[str]] = None) -> Dict[str, bool]:
        """Add all files in a directory to the knowledge base."""
        if file_extensions is None:
            file_extensions = ['.pdf', '.docx', '.xlsx', '.xls', '.csv', '.txt', '.html', '.htm']

        directory = Path(directory_path)
        if not directory.is_dir():
            logger.error(f"{directory_path} is not a valid directory")
            return {}

        # Collect files to process
        files_to_process = []

        if recursive:
            for ext in file_extensions:
                files_to_process.extend(list(directory.rglob(f"*{ext}")))
        else:
            for ext in file_extensions:
                files_to_process.extend(list(directory.glob(f"*{ext}")))

        logger.info(f"Found {len(files_to_process)} files to process in {directory_path}")

        # Process files in parallel
        return self.batch_add_files([str(f) for f in files_to_process])

    def _expand_query(self, query: str) -> List[str]:
        """Expand query with related terms to improve recall."""
        # Start with original query
        expanded_queries = [query]

        # Skip expansion for very short queries
        if len(query.split()) <= 2:
            return expanded_queries

        try:
            # 1. Add a variant with different word order for multi-word queries
            words = query.split()
            if len(words) >= 3:
                # Shuffle words a bit - move first word to end
                reordered = ' '.join(words[1:] + [words[0]])
                expanded_queries.append(reordered)

            # 2. Add synonyms for key terms
            try:
                from nltk.corpus import wordnet

                # Try to add WordNet synonyms
                for word in words:
                    if len(word) <= 3 or word.lower() in self.text_processor.stopwords:
                        continue  # Skip short words and stopwords

                    synonyms = []
                    for syn in wordnet.synsets(word):
                        for lemma in syn.lemmas():
                            synonym = lemma.name().replace('_', ' ')
                            if synonym != word and len(synonym) > 3:
                                synonyms.append(synonym)

                    # Add up to 2 synonyms for each word
                    for synonym in synonyms[:2]:
                        new_query = query.replace(word, synonym)
                        if new_query != query:
                            expanded_queries.append(new_query)

                    # If we have enough expansions, stop
                    if len(expanded_queries) >= 5:
                        break
            except Exception:
                pass

            # 3. Add abbreviation/expansion variants
            abbrev_expansions = self._detect_abbreviations(query)
            for expanded in abbrev_expansions:
                if expanded != query and expanded not in expanded_queries:
                    expanded_queries.append(expanded)

            # 4. Remove some words for broader matching
            if len(words) >= 4:
                shortened = ' '.join(words[:-1])  # Remove last word
                expanded_queries.append(shortened)

            # Limit the number of expanded queries
            return expanded_queries[:5]  # Return at most 5 queries

        except Exception as e:
            logger.warning(f"Query expansion error: {e}")
            return [query]  # Return just the original query

    def _detect_abbreviations(self, text: str) -> List[str]:
        """Detect and expand abbreviations in text."""
        results = []

        # Common abbreviations in technical documents
        common_abbrevs = {
            "pdf": "portable document format",
            "doc": "document",
            "xlsx": "excel spreadsheet",
            "api": "application programming interface",
            "ui": "user interface",
            "db": "database",
            "kb": "knowledge base",
            "ai": "artificial intelligence",
            "ml": "machine learning",
            "nlp": "natural language processing",
            "poc": "proof of concept",
            "qa": "quality assurance",
            "sme": "subject matter expert",
            "sow": "statement of work",
            "rfp": "request for proposal",
        }

        # Look for common abbreviations
        for abbrev, expansion in common_abbrevs.items():
            if f" {abbrev} " in f" {text.lower()} ":
                # Replace abbreviation with expansion
                expanded = re.sub(r'\b' + re.escape(abbrev) + r'\b', expansion, text.lower(), flags=re.IGNORECASE)
                if expanded != text.lower():
                    results.append(expanded)

        return results

    def _cross_encoder_rerank(self, query: str, candidate_texts: List[str]) -> Optional[np.ndarray]:
        """Rerank results using a cross-encoder model if available."""
        try:
            from sentence_transformers import CrossEncoder

            # Use a cross-encoder model for more accurate relevance scoring
            # First check if we can access a cross-encoder model
            model_name = "cross-encoder/ms-marco-MiniLM-L-6-v2"

            model = CrossEncoder(model_name, max_length=512)

            # Prepare text pairs for scoring
            text_pairs = [(query, text) for text in candidate_texts]

            # Get relevance scores
            scores = model.predict(text_pairs)

            # Normalize scores
            if len(scores) > 1:
                min_score = min(scores)
                max_score = max(scores)
                if max_score > min_score:
                    norm_scores = (scores - min_score) / (max_score - min_score)
                else:
                    norm_scores = np.ones_like(scores)
                return norm_scores
            else:
                return np.ones_like(scores)

        except Exception as e:
            logger.warning(f"Cross-encoder reranking not available: {e}")
            return None

    def get_stats(self) -> Dict[str, Any]:
        """Get enhanced statistics about the knowledge base."""
        try:
            # Basic stats
            stats = {
                "total_chunks": len(self.chunks),
                "total_files": len(self.file_registry),
                "index_dimension": self.dimension,
                "embedding_model": self.embedding_model_name,
                "embedding_provider": self.embedding_provider.value,
                "chunking_strategy": self.chunking_strategy.value,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "file_types": {},
                "last_modified": datetime.fromtimestamp(os.path.getmtime(self.metadata_path)).isoformat()
                    if self.metadata_path.exists() else None,
                "total_tokens": sum(len(chunk.split()) for chunk in self.chunks),
            }

            # Count file types
            source_types = {}
            for metadata in self.metadatas:
                source_type = metadata.source_type
                if source_type in source_types:
                    source_types[source_type] += 1
                else:
                    source_types[source_type] = 1

            stats["source_types"] = source_types

            # Count file extensions
            for file_name in self.file_registry:
                ext = os.path.splitext(file_name)[1].lower()
                if ext in stats["file_types"]:
                    stats["file_types"][ext] += 1
                else:
                    stats["file_types"][ext] = 1

            # Enhanced stats

            # 1. Content language distribution
            language_dist = {}
            sample_chunks = random.sample(self.chunks, min(100, len(self.chunks))) if self.chunks else []
            for chunk in sample_chunks:
                lang = self.text_processor.detect_language(chunk)
                language_dist[lang] = language_dist.get(lang, 0) + 1
            stats["language_distribution"] = language_dist

            # 2. Recent activity
            now = int(time.time())
            one_week_ago = now - (7 * 24 * 60 * 60)
            recent_files = {}
            for file_name, info in self.file_registry.items():
                if info.get("timestamp", 0) >= one_week_ago:
                    recent_files[file_name] = info.get("timestamp")
            stats["recent_activity"] = {
                "files_added_last_7_days": len(recent_files),
                "most_recent_file": max(recent_files.items(), key=lambda x: x[1])[0] if recent_files else None
            }

            # 3. Average chunk sizes
            chunk_sizes = [len(chunk.split()) for chunk in self.chunks]
            if chunk_sizes:
                stats["chunk_stats"] = {
                    "avg_tokens_per_chunk": sum(chunk_sizes) / len(chunk_sizes),
                    "min_chunk_size": min(chunk_sizes),
                    "max_chunk_size": max(chunk_sizes),
                }

            # 4. Table statistics
            table_count = sum(1 for meta in self.metadatas if meta.is_table)
            stats["tables"] = {
                "total_tables": table_count,
                "table_percentage": (table_count / len(self.metadatas)) * 100 if self.metadatas else 0
            }

            # 5. Storage statistics
            if self.index_path.exists() and self.metadata_path.exists():
                index_size = os.path.getsize(self.index_path)
                metadata_size = os.path.getsize(self.metadata_path)
                stats["storage"] = {
                    "index_size_mb": round(index_size / (1024 * 1024), 2),
                    "metadata_size_mb": round(metadata_size / (1024 * 1024), 2),
                    "total_size_mb": round((index_size + metadata_size) / (1024 * 1024), 2)
                }

            return stats
        except Exception as e:
            logger.warning(f"Error generating extended stats: {e}")
            # Fall back to basic stats
            return {
                "total_chunks": len(self.chunks),
                "total_files": len(self.file_registry),
                "embedding_model": self.embedding_model_name
            }